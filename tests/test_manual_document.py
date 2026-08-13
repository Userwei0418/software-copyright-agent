import hashlib
import json
import tempfile
import unittest
import zipfile
from pathlib import Path
from uuid import uuid4

from docx import Document
from PIL import Image, ImageDraw

from software_copyright_agent.manual_document import ManualDocumentService
from software_copyright_agent.manual_pipeline import ManualPipelineService
from software_copyright_agent.storage import Database


class ManualDocumentServiceTests(unittest.TestCase):
    def test_internal_screenshot_workflow_text_and_generic_heading_never_ship(self) -> None:
        service = ManualDocumentService.__new__(ManualDocumentService)
        self.assertTrue(service._internal_ui_delivery_block({
            "type": "paragraph",
            "text": "本章节等待用户提供并确认真实截图证据。其他正文章节可继续审阅与导出。",
        }))
        self.assertTrue(service._internal_ui_delivery_block({
            "type": "subheading", "title": "截屏",
        }))
        self.assertFalse(service._internal_ui_delivery_block({
            "type": "subheading", "title": "首页与导航",
        }))

    def test_assembles_native_word_sections_tables_figures_and_screenshots(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            database = Database(root / "app.db")
            database.initialize()
            now = "2026-08-10T00:00:00Z"
            with database.connect() as connection:
                connection.execute(
                    """INSERT INTO project_sources(id,kind,original_path,display_name,
                    created_at,last_opened_at) VALUES
                    ('source','directory','/tmp/project','证据化系统',?,?)""", (now, now),
                )
                connection.execute(
                    """INSERT INTO project_snapshots(id,source_id,root_fingerprint,
                    scanner_version,rules_version,summary_json,manifest_relative_path,created_at)
                    VALUES ('snapshot','source','hash','v1','v1','{}','manifest.jsonl',?)""",
                    (now,),
                )
                connection.execute(
                    """INSERT INTO tasks(id,source_id,snapshot_id,status,workflow_version,
                    quality_policy_version,created_at,updated_at) VALUES
                    ('task','source','snapshot','completed','v1','v1',?,?)""", (now, now),
                )
                for key, value in (("project.name", "证据化系统"),
                                   ("project.version", "V2.0")):
                    connection.execute(
                        """INSERT INTO facts(id,task_id,fact_key,value_json,status,source,
                        confidence,evidence_ids_json,created_at,confirmed_at) VALUES
                        (?,?,?,?,'confirmed','user',1,'[]',?,?)""",
                        (str(uuid4()), "task", key, json.dumps(value), now, now),
                    )
                connection.execute(
                    """INSERT INTO model_configs(id,name,protocol_id,base_url,model_name,
                    settings_json,enabled,created_at,updated_at) VALUES
                    ('model-config','Local','ollama','http://127.0.0.1:11434','model',
                    '{}',1,?,?)""", (now, now),
                )
            job = ManualPipelineService(database).create("task", "model-config")
            blocks = [
                {"type": "subheading", "title": "证据处理与产物追踪",
                 "evidence_refs": ["ref"], "inference": False},
                {"type": "paragraph", "text": "本系统围绕证据化处理组织业务能力，所有阶段结果均在本地持久化并支持恢复。",
                 "evidence_refs": ["ref"], "inference": False},
                {"type": "list", "lead": "核心职责包括：", "items": ["读取并校验项目材料", "保存可追溯的生成结果"],
                 "evidence_refs": ["ref"], "inference": False},
                {"type": "table", "title": "模块职责", "headers": ["模块", "职责"],
                 "rows": [["扫描模块", "提取项目事实"], ["文档模块", "装配正式说明书"]],
                 "evidence_refs": ["ref"], "inference": False},
            ]
            figure_blocks = blocks + [{
                "type": "figure_request", "figure_key": "architecture",
                "figure_type": "architecture", "title": "系统架构图",
                "purpose": "展示模块协作", "evidence_refs": ["ref"],
            }]
            with database.connect() as connection:
                for ordinal, (key, title, content) in enumerate((
                    ("introduction", "引言", blocks),
                    ("architecture", "总体设计", figure_blocks),
                    ("modules", "功能与模块设计", blocks),
                    ("ui_operations", "用户界面与操作说明", blocks),
                ), 1):
                    if key == "ui_operations":
                        ordinal = 7
                    connection.execute(
                        """INSERT INTO manual_section_artifacts(id,job_id,section_key,title,
                        ordinal,status,content_json,evidence_refs_json,inference_notes_json,
                        figure_requests_json,updated_at) VALUES
                        (?,?,?,?,?,'generated',?,'[\"ref\"]','[]','[]',?)""",
                        (str(uuid4()), job["id"], key, title, ordinal,
                         json.dumps(content, ensure_ascii=False), now),
                    )
            task_root = root / "tasks" / "task"
            service = ManualDocumentService(database, root)
            checkpoint = service.assemble_checkpoint(job["id"])
            self.assertEqual(checkpoint["version"], 1)
            self.assertEqual(checkpoint["document_kind"], "review_checkpoint")
            self.assertEqual(
                checkpoint["filename"], "证据化系统-V2.0-软件说明书-阶段审阅稿.docx"
            )
            self.assertEqual(checkpoint["qa"]["figure_count"], 0)
            self.assertEqual(checkpoint["qa"]["screenshot_count"], 0)
            checkpoint_doc = Document(
                task_root / checkpoint["docx_relative_path"]
            )
            checkpoint_text = "\n".join(item.text for item in checkpoint_doc.paragraphs)
            self.assertIn("软件说明书 · 阶段审阅稿", checkpoint_text)
            self.assertIn("图表、界面截图与最终质量检查尚未完成", checkpoint_text)
            checkpoint_job = ManualPipelineService(database).get(job["id"])
            self.assertEqual(checkpoint_job["current_step"], "research")
            self.assertEqual(checkpoint_job["progress"]["completed"], 0)
            figure_relative = "artifacts/manual-figures/architecture.v1.png"
            planned_figure_relative = "artifacts/manual-figures/modules.v1.png"
            screenshot_relative = "artifacts/manual-screenshots/main.v1.png"
            for relative, label in ((figure_relative, "ARCH"),
                                    (planned_figure_relative, "MODULES"),
                                    (screenshot_relative, "UI")):
                path = task_root / relative
                path.parent.mkdir(parents=True, exist_ok=True)
                image = Image.new("RGB", (1200, 700), "white")
                ImageDraw.Draw(image).text((80, 80), label, fill="black")
                image.save(path, "PNG")
            with database.connect() as connection:
                connection.execute(
                    """INSERT INTO manual_figure_artifacts(id,job_id,figure_key,section_key,
                    figure_type,title,status,semantic_json,drawio_relative_path,svg_relative_path,
                    png_relative_path,qa_json,updated_at) VALUES
                    (?,?, 'architecture','architecture','architecture','系统架构图','verified',
                    '{}','a.drawio','a.svg',?,'{}',?)""",
                    (str(uuid4()), job["id"], figure_relative, now),
                )
                # The deterministic planner may add a useful chapter figure even
                # when the AI正文 omitted a figure_request placeholder.
                connection.execute(
                    """INSERT INTO manual_figure_artifacts(id,job_id,figure_key,section_key,
                    figure_type,title,status,semantic_json,drawio_relative_path,svg_relative_path,
                    png_relative_path,qa_json,updated_at) VALUES
                    (?,?,'module-collaboration','modules','module','模块协作图','verified',
                    '{}','m.drawio','m.svg',?,'{}',?)""",
                    (str(uuid4()), job["id"], planned_figure_relative, now),
                )
                interpretation = {
                    "page_title": "项目工作台", "page_type": "dashboard",
                    "purpose": "用于查看当前项目的整体处理状态和主要材料入口。",
                    "target_roles": ["材料编制人员"],
                    "entry_conditions": ["用户完成项目扫描并从项目列表进入当前任务。"],
                    "visible_regions": ["导航区", "状态区", "操作区", "结果预览区"],
                    "key_controls": ["生成说明书"],
                    "workflow_steps": ["确认项目", "启动生成", "查看文档结果"],
                    "success_state": "页面展示生成结果。",
                    "failure_and_recovery": "失败时保留阶段记录并允许重试。",
                    "related_backend_actions": ["读取本地任务与产物状态"],
                    "route_guess": "/manual", "related_evidence_refs": ["ref"],
                    "suggested_group": "项目管理", "suggested_order": 1,
                    "suggested_caption": "项目工作台页面", "confidence": 0.95,
                    "warnings": [],
                }
                connection.execute(
                    """INSERT INTO manual_project_profile_revisions(
                    id,task_id,version,research_artifact_id,origin,profile_json,fingerprint,created_at)
                    VALUES('profile','task',1,NULL,'user','{}','profile-hash',?)""", (now,),
                )
                connection.execute(
                    """INSERT INTO manual_project_screenshot_assets(
                    id,task_id,asset_key,source,title,image_relative_path,width,height,image_format,
                    sha256,analysis_status,review_status,adoption_status,group_key,group_title,
                    sort_order,sensitive_status,created_at,updated_at)
                    VALUES('asset','task','main','user','项目工作台',?,1200,700,'PNG','image-hash',
                    'completed','reviewed','adopted','screenshots','截屏',1,'confirmed_safe',?,?)""",
                    (screenshot_relative, now, now),
                )
                connection.execute(
                    """INSERT INTO manual_project_screenshot_revisions(
                    id,asset_id,version,title,image_relative_path,width,height,image_format,sha256,
                    edit_source,parent_revision_id,created_at)
                    VALUES('asset-r1','asset',1,'项目工作台',?,1200,700,'PNG','image-hash',
                    'import',NULL,?)""", (screenshot_relative, now),
                )
                connection.execute(
                    """INSERT INTO manual_screenshot_interpretation_revisions(
                    id,asset_id,asset_revision_id,project_profile_revision_id,version,
                    model_config_id,model_name,prompt_version,cache_key,status,
                    interpretation_json,origin,reviewed,attempt_count,elapsed_ms,created_at)
                    VALUES('interpretation-r1','asset','asset-r1','profile',1,NULL,'人工审核',
                    'test','cache','completed',?,'user',1,1,0,?)""",
                    (json.dumps(interpretation, ensure_ascii=False), now),
                )
                connection.execute(
                    """INSERT INTO manual_job_screenshot_refs(
                    id,job_id,asset_id,asset_revision_id,interpretation_revision_id,
                    group_key,group_title,sort_order,adopted_at)
                    VALUES(?,?,'asset','asset-r1','interpretation-r1','screenshots','截屏',1,?)""",
                    (str(uuid4()), job["id"], now),
                )

            result = service.assemble(job["id"])
            self.assertEqual(result["version"], 2)
            self.assertEqual(result["document_kind"], "formal_candidate")
            self.assertEqual(result["filename"], "证据化系统-V2.0-软件说明书-审阅稿.docx")
            self.assertEqual(result["integrity"]["status"], "verified")
            self.assertEqual(result["qa"]["figure_count"], 2)
            self.assertEqual(result["qa"]["screenshot_count"], 1)
            self.assertTrue(result["docx_relative_path"].startswith(
                "artifacts/manual/jobs/job-v1/"))
            artifact = task_root / result["docx_relative_path"]
            doc = Document(artifact)
            self.assertAlmostEqual(
                doc.styles["Normal"].paragraph_format.line_spacing, 1.18, places=2
            )
            self.assertAlmostEqual(
                doc.styles["Normal"].paragraph_format.space_after.pt, 4.0
            )
            self.assertGreaterEqual(len(doc.inline_shapes), 3)
            self.assertEqual(len(doc.tables), 3)
            captions = [
                paragraph.text for paragraph in doc.paragraphs
                if paragraph.style.name == "Caption"
            ]
            self.assertIn("图 1  系统架构图", captions)
            self.assertIn("图 2  模块协作图", captions)
            self.assertIn("图 3  项目工作台页面", captions)
            self.assertEqual(
                [value for value in captions if value.startswith("表 ")],
                ["表 1  模块职责", "表 2  模块职责", "表 3  模块职责"],
            )
            headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
            self.assertIn("1  引言", headings)
            self.assertIn("3  功能与模块设计", headings)
            self.assertIn("7  用户界面与操作说明", headings)
            screenshot_headings = [
                p for p in doc.paragraphs if p.style.name == "Heading 2"
                and p.text == "7.1  项目工作台"
            ]
            self.assertEqual(len(screenshot_headings), 1)
            self.assertIn("3.1  证据处理与产物追踪", [
                p.text for p in doc.paragraphs if p.style.name == "Heading 2"
            ])
            ui_headings = [
                p.text for p in doc.paragraphs if p.style.name == "Heading 2"
                and p.text.startswith("7.")
            ]
            self.assertEqual(ui_headings, ["7.1  项目工作台"])
            self.assertFalse(screenshot_headings[0].paragraph_format.page_break_before)
            self.assertIn("对应操作界面见图 3", "\n".join(p.text for p in doc.paragraphs))
            self.assertIn("该功能模块主要用于", "\n".join(p.text for p in doc.paragraphs))
            with zipfile.ZipFile(artifact) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
                settings_xml = archive.read("word/settings.xml").decode("utf-8")
                font_table_xml = archive.read("word/fontTable.xml").decode("utf-8")
                font_relationships_xml = archive.read(
                    "word/_rels/fontTable.xml.rels"
                ).decode("utf-8")
                embedded_font = archive.read("word/fonts/NotoSansCJKsc-Regular.odttf")
            self.assertIn("章节导航", document_xml)
            self.assertIn("01    引言", document_xml)
            self.assertIn("PAGEREF manual_section_1", document_xml)
            self.assertIn('w:name="manual_section_1"', document_xml)
            self.assertIn('w:numFmt w:val="bullet"', numbering_xml)
            self.assertIn('w:updateFields w:val="true"', settings_xml)
            self.assertIn('w:name="Noto Sans CJK SC"', font_table_xml)
            self.assertIn("w:embedRegular", font_table_xml)
            self.assertIn("relationships/font", font_relationships_xml)
            self.assertGreater(len(embedded_font), 1_000_000)
            before_toc_sha = result["sha256"]
            updated_toc = service.persist_toc_page_results(
                job["id"], 2, {1: 3, 2: 4, 3: 5}
            )
            self.assertNotEqual(updated_toc["sha256"], before_toc_sha)
            toc_text = "\n".join(
                paragraph.text for paragraph in Document(artifact).paragraphs
            )
            self.assertIn("01    引言\t3", toc_text)
            self.assertIn("03    功能与模块设计\t5", toc_text)
            self.assertEqual(service.read(job["id"], 2), artifact.read_bytes())
            self.assertEqual(result["freshness"]["status"], "current")
            with database.connect() as connection:
                connection.execute(
                    "UPDATE manual_document_artifacts SET status='qa_passed' WHERE id=?",
                    (result["id"],),
                )
                connection.execute(
                    """INSERT INTO manual_document_qa_runs(id,document_artifact_id,job_id,
                    qa_version,policy_version,renderer_kind,passed,checks_json,summary_json,
                    report_relative_path,render_relative_path,preview_pdf_relative_path,created_at)
                    VALUES (?,?,?,1,'manual-docx-qa-v15','libreoffice_word',1,
                    '[]','{}','report-current.json','render-current','preview-current.pdf',?)""",
                    (str(uuid4()), result["id"], job["id"], now),
                )
            final = service.finalize(job["id"], 2)
            self.assertEqual(final["document_kind"], "final_document")
            self.assertEqual(final["filename"], "证据化系统-V2.0-软件说明书.docx")
            final_text = "\n".join(
                paragraph.text for paragraph in Document(
                    task_root / final["docx_relative_path"]
                ).paragraphs
            )
            self.assertNotIn("章节导航（页码将在 Word 打开或导出时自动更新）", final_text)
            self.assertNotIn("图中内容以已审核的真实截图为准", final_text)
            self.assertNotIn("页面实际内容和操作范围以图中已审核的真实界面为准", final_text)
            self.assertIn("软件说明书", final_text)
            self.assertNotIn("软件说明书 · 终稿", final_text)
            final_doc = Document(task_root / final["docx_relative_path"])
            self.assertEqual(str(final_doc.styles["Heading 1"].font.color.rgb), "172331")
            self.assertEqual(str(final_doc.styles["Heading 2"].font.color.rgb), "172331")
            detail_style = final_doc.styles["Manual Screenshot Detail"]
            self.assertAlmostEqual(detail_style.font.size.pt, 10.0)
            self.assertAlmostEqual(detail_style.paragraph_format.left_indent.pt, 21.0)
            with database.connect() as connection:
                connection.execute(
                    "DELETE FROM manual_document_qa_runs WHERE document_artifact_id=?",
                    (result["id"],),
                )
                connection.execute(
                    "UPDATE manual_document_artifacts SET status='qa_passed' WHERE id=?",
                    (result["id"],),
                )
                connection.execute(
                    """INSERT INTO manual_document_qa_runs(id,document_artifact_id,job_id,
                    qa_version,policy_version,renderer_kind,passed,checks_json,summary_json,
                    report_relative_path,render_relative_path,preview_pdf_relative_path,created_at)
                    VALUES (?,?,?,1,'manual-docx-qa-v1','deterministic_companion',1,
                    '[]','{}','report.json','render','preview.pdf',?)""",
                    (str(uuid4()), result["id"], job["id"], now),
                )
            historical = service.get(job["id"], 2)
            self.assertEqual(historical["quality"]["status"], "outdated")
            self.assertFalse(historical["quality"]["current_policy"])
            self.assertTrue(historical["quality"]["current_generator"])
            with database.connect() as connection:
                connection.execute(
                    "UPDATE manual_figure_artifacts SET updated_at='2099-01-01T00:00:00Z' WHERE job_id=?",
                    (job["id"],),
                )
            self.assertEqual(service.list(job["id"])[0]["freshness"]["status"], "outdated")
            with database.connect() as connection:
                connection.execute(
                    "UPDATE manual_figure_artifacts SET updated_at=? WHERE job_id=?",
                    (now, job["id"]),
                )
                connection.execute(
                    "UPDATE manual_project_screenshot_assets SET updated_at='2099-01-02T00:00:00Z' WHERE id='asset'",
                )
            screenshot_freshness = service.list(job["id"])[0]["freshness"]
            self.assertEqual(screenshot_freshness["status"], "outdated")
            pipeline = ManualPipelineService(database).get(job["id"])
            self.assertEqual(pipeline["current_step"], "render_qa")
            self.assertEqual(pipeline["progress"]["completed"], 5)


if __name__ == "__main__":
    unittest.main()

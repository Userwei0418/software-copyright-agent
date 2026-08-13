import hashlib
import json
import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional
from uuid import uuid4

from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageChops, ImageDraw, ImageFont

from .font_assets import FontAsset
from .manual_document import ManualDocumentError, ManualDocumentService
from .manual_drafting import unverified_outcome_hits
from .service import utc_now
from .source_document_qa import LibreOfficeRenderer, SourceDocumentQaError
from .storage import Database


QA_POLICY_VERSION = "manual-docx-qa-v16"
# Default used by the deterministic renderer in focused unit tests. Production
# persists the concrete runtime renderer so previews and QA evidence agree.
RENDERER_KIND = "deterministic_companion"
A4_PIXELS = (1191, 1684)
INK = "#172331"
ACCENT = INK
ACCENT_DARK = INK
MUTED = INK
TABLE_FILL = "#e8eef5"


class ManualQaError(ValueError):
    pass


@dataclass(frozen=True)
class ManualQaCheck:
    key: str
    passed: bool
    severity: str
    expected: object
    actual: object
    message: str


@dataclass(frozen=True)
class CompanionRenderResult:
    pdf_path: Path
    page_paths: tuple
    page_kinds: tuple
    fill_ratios: tuple
    underfilled_pages: tuple


@dataclass(frozen=True)
class ManualQaResult:
    passed: bool
    checks: tuple
    summary: dict
    render: CompanionRenderResult


class _Page:
    def __init__(self, number: int, kind: str, header: str, font, muted_font) -> None:
        self.number = number
        self.kind = kind
        self.image = Image.new("RGB", A4_PIXELS, "white")
        self.draw = ImageDraw.Draw(self.image)
        self.y = 138
        self.min_y = A4_PIXELS[1]
        self.max_y = 0
        if kind != "cover":
            self.draw.text((1060, 65), header, font=muted_font, fill=MUTED, anchor="ra")
        footer = "第 {0} 页".format(number)
        self.draw.text((A4_PIXELS[0] // 2, 1616), footer, font=font, fill=MUTED, anchor="ma")

    def mark(self, top: int, bottom: int) -> None:
        self.min_y = min(self.min_y, top)
        self.max_y = max(self.max_y, bottom)


class ManualCompanionRenderer:
    """Renders the same structured source as the DOCX without office dependencies."""

    renderer_kind = "deterministic_companion"

    left = 126
    right = 1065
    content_bottom = 1532

    def __init__(self, font_asset: FontAsset = None) -> None:
        self.font_asset = font_asset or FontAsset.bundled_cjk()
        path = str(self.font_asset.path)
        self.fonts = {
            "cover_kicker": ImageFont.truetype(path, 23),
            "cover_title": ImageFont.truetype(path, 55),
            "cover_subtitle": ImageFont.truetype(path, 39),
            "cover_version": ImageFont.truetype(path, 27),
            "h1": ImageFont.truetype(path, 33),
            "h2": ImageFont.truetype(path, 27),
            "body": ImageFont.truetype(path, 22),
            "body_bold": ImageFont.truetype(path, 22),
            "small": ImageFont.truetype(path, 18),
            "caption": ImageFont.truetype(path, 18),
            "table": ImageFont.truetype(path, 19),
        }
        self.pages = []
        self.context = {}

    def render(self, document_path: Path, output_dir: Path, context: dict,
               sections: list, figures: list,
               screenshots: list, task_root: Path) -> CompanionRenderResult:
        output_dir.mkdir(parents=True, exist_ok=True)
        self.pages = []
        self.context = context
        self._cover()
        self._toc(sections)
        figure_map = {item["figure_key"]: item for item in figures}
        screenshots_by_section = {}
        for item in screenshots:
            screenshots_by_section.setdefault(item["section_key"], []).append(item)
        self._new_page("content")
        for index, section in enumerate(sections, 1):
            self._heading("{0}  {1}".format(index, section["title"]))
            subsection_index = 0
            for block in section["blocks"]:
                kind = block.get("type")
                if kind == "subheading":
                    subsection_index += 1
                    self._subheading("{0}.{1}  {2}".format(
                        index, subsection_index, block.get("title", "")
                    ))
                elif kind == "paragraph":
                    self._paragraph(str(block.get("text", "")))
                elif kind == "list":
                    lead = str(block.get("lead", "")).strip()
                    if lead:
                        self._paragraph(lead)
                    for item in block.get("items", []):
                        self._bullet(str(item))
                elif kind == "table":
                    self._table(block)
                elif kind == "figure_request":
                    figure = figure_map.get(block.get("figure_key"))
                    if figure and figure.get("png_relative_path"):
                        path = self._asset(task_root, figure["png_relative_path"])
                        if path:
                            self._picture(path, "图  " + figure["title"], 510)
            for screenshot in screenshots_by_section.get(section["section_key"], []):
                path = self._asset(task_root, screenshot["image_relative_path"])
                if path:
                    self._screenshot(path, screenshot["title"], screenshot["description"])
        if self.pages and self.pages[-1].kind == "content" and self.pages[-1].max_y == 0:
            self.pages.pop()
        page_paths = []
        for page in self.pages:
            path = output_dir / "page-{0}.png".format(page.number)
            page.image.save(path, "PNG", optimize=True)
            page_paths.append(path)
        if not page_paths:
            raise ManualQaError("内置渲染器没有生成页面")
        pdf_path = output_dir / "preview.pdf"
        first, remaining = self.pages[0].image, [page.image for page in self.pages[1:]]
        first.save(pdf_path, "PDF", resolution=144, save_all=True, append_images=remaining)
        ratios = []
        underfilled = []
        usable = self.content_bottom - 138
        for page in self.pages:
            if page.kind != "content":
                ratios.append(1.0)
                continue
            ratio = max(0.0, min(1.0, (page.max_y - 138) / usable))
            ratios.append(round(ratio, 4))
            if ratio < 0.35:
                underfilled.append(page.number)
        return CompanionRenderResult(
            pdf_path, tuple(page_paths), tuple(page.kind for page in self.pages),
            tuple(ratios), tuple(underfilled),
        )

    def _new_page(self, kind: str) -> _Page:
        page = _Page(
            len(self.pages) + 1, kind,
            "{0} · 软件说明书".format(self.context.get("project_name", "")),
            self.fonts["small"], self.fonts["small"],
        )
        self.pages.append(page)
        return page

    @property
    def page(self) -> _Page:
        return self.pages[-1]

    def _cover(self) -> None:
        page = self._new_page("cover")
        center = A4_PIXELS[0] // 2
        items = (
            (490, "软件著作权登记材料", self.fonts["cover_kicker"], ACCENT),
            (590, str(self.context["project_name"]), self.fonts["cover_title"], INK),
            (710, "软件说明书", self.fonts["cover_subtitle"], ACCENT_DARK),
            (825, str(self.context["project_version"]), self.fonts["cover_version"], MUTED),
            (1045, "技术设计与使用说明", self.fonts["body"], MUTED),
        )
        for y, text, font, color in items:
            page.draw.text((center, y), text, font=font, fill=color, anchor="mm",
                           stroke_width=1 if font in {self.fonts["cover_title"], self.fonts["cover_subtitle"]} else 0)
            box = page.draw.textbbox((center, y), text, font=font, anchor="mm")
            page.mark(box[1], box[3])

    def _toc(self, sections: list) -> None:
        page = self._new_page("toc")
        self._draw_text(page, self.left, 175, "目录", self.fonts["h1"], ACCENT, bold=True)
        self._draw_text(page, self.left, 255, "章节目录", self.fonts["body"], MUTED)
        y = 345
        for index, section in enumerate(sections, 1):
            self._draw_text(page, self.left + 20, y, "{0:02d}    {1}".format(index, section["title"]),
                            self.fonts["body"], INK, bold=True)
            y += 72

    def _ensure(self, required: int) -> None:
        if self.page.kind != "content" or self.page.y + required > self.content_bottom:
            self._new_page("content")

    def _heading(self, text: str) -> None:
        self._ensure(78)
        page, top = self.page, self.page.y
        self._draw_text(page, self.left, top, text, self.fonts["h1"], ACCENT, bold=True)
        page.y += 70

    def _subheading(self, text: str) -> None:
        self._ensure(62)
        page, top = self.page, self.page.y
        self._draw_text(page, self.left, top, text, self.fonts["h2"], ACCENT_DARK, bold=True)
        page.y += 54
        page.mark(top, page.y)

    def _paragraph(self, text: str) -> None:
        text = text.strip()
        if not text:
            return
        lines = self._wrap(text, self.fonts["body"], self.right - self.left - 50)
        line_height = 36
        self._ensure(len(lines) * line_height + 18)
        page, top = self.page, self.page.y
        for index, line in enumerate(lines):
            x = self.left + (44 if index == 0 else 0)
            self._draw_text(page, x, page.y, line, self.fonts["body"], INK)
            page.y += line_height
        page.y += 13
        page.mark(top, page.y)

    def _bullet(self, text: str) -> None:
        lines = self._wrap(text.strip(), self.fonts["body"], self.right - self.left - 72)
        line_height = 35
        self._ensure(max(1, len(lines)) * line_height + 8)
        page, top = self.page, self.page.y
        self._draw_text(page, self.left + 26, page.y + 1, "•", self.fonts["body"], INK)
        for line in lines:
            self._draw_text(page, self.left + 62, page.y, line, self.fonts["body"], INK)
            page.y += line_height
        page.y += 6
        page.mark(top, page.y)

    def _table(self, block: dict) -> None:
        headers = [str(item) for item in block.get("headers", [])]
        rows = [[str(cell) for cell in row] for row in block.get("rows", [])]
        if not headers or not rows:
            return
        columns = len(headers)
        weights = []
        for index, header in enumerate(headers):
            values = [header] + [row[index] for row in rows if index < len(row)]
            weights.append(min(32, max(7, max(len(value) for value in values))))
        total = sum(weights)
        widths = [round((self.right - self.left) * value / total) for value in weights]
        widths[-1] += self.right - self.left - sum(widths)
        prepared = []
        for row in [headers] + rows:
            cells = []
            line_count = 1
            for index in range(columns):
                value = row[index] if index < len(row) else ""
                lines = self._wrap(value, self.fonts["table"], widths[index] - 24)
                cells.append(lines)
                line_count = max(line_count, len(lines))
            prepared.append((cells, max(52, line_count * 29 + 18)))
        caption_height = 45 if block.get("title") else 10
        total_height = caption_height + sum(height for _, height in prepared) + 22
        if total_height < self.content_bottom - 138:
            self._ensure(total_height)
        if block.get("title"):
            self._ensure(45 + prepared[0][1])
            caption = "表  {0}".format(block["title"])
            width = self.page.draw.textlength(caption, font=self.fonts["caption"])
            self._draw_text(self.page, (A4_PIXELS[0] - width) // 2, self.page.y,
                            caption, self.fonts["caption"], MUTED)
            self.page.y += 42
        header = prepared[0]
        for row_index, (cells, row_height) in enumerate(prepared):
            if self.page.y + row_height > self.content_bottom:
                self._new_page("content")
                if row_index > 0:
                    self._draw_table_row(header[0], header[1], widths, True)
            self._draw_table_row(cells, row_height, widths, row_index == 0)
        self.page.y += 20

    def _draw_table_row(self, cells: list, height: int, widths: list, header: bool) -> None:
        page, top, x = self.page, self.page.y, self.left
        for index, lines in enumerate(cells):
            width = widths[index]
            page.draw.rectangle((x, top, x + width, top + height),
                                fill=TABLE_FILL if header else "white", outline="#3d4952", width=1)
            line_height = 29
            y = top + (height - len(lines) * line_height) // 2
            for line in lines:
                text_width = page.draw.textlength(line, font=self.fonts["table"])
                tx = x + max(12, (width - text_width) / 2)
                self._draw_text(page, tx, y, line, self.fonts["table"], INK, bold=header)
                y += line_height
            x += width
        page.y += height
        page.mark(top, page.y)

    def _picture(self, path: Path, caption: str, max_height: int) -> None:
        with Image.open(path) as source:
            image = source.convert("RGB")
        max_width = self.right - self.left - 24
        scale = min(max_width / image.width, max_height / image.height, 1.0)
        size = (max(1, round(image.width * scale)), max(1, round(image.height * scale)))
        image = image.resize(size, Image.Resampling.LANCZOS)
        required = image.height + 65
        self._ensure(required)
        page, top = self.page, self.page.y
        x = (A4_PIXELS[0] - image.width) // 2
        page.image.paste(image, (x, page.y))
        page.y += image.height + 18
        text_width = page.draw.textlength(caption, font=self.fonts["caption"])
        self._draw_text(page, (A4_PIXELS[0] - text_width) / 2, page.y,
                        caption, self.fonts["caption"], MUTED)
        page.y += 42
        page.mark(top, page.y)

    def _screenshot(self, path: Path, title: str, description: dict) -> None:
        labels = (
            ("page_purpose", "页面用途"), ("entry_conditions", "进入条件"),
            ("visible_regions", "可见区域与控件"), ("typical_workflow", "典型操作流程"),
            ("backend_interactions", "后台、接口与数据交互"),
            ("result_validation_recovery", "结果、校验与异常恢复"),
        )
        descriptions = []
        estimated = 0
        for key, label in labels:
            value = str(description.get(key, "")).strip()
            if not value:
                continue
            lines = self._wrap(label + "：" + value, self.fonts["small"], self.right - self.left - 34)
            descriptions.append((label, value, lines))
            estimated += len(lines) * 30 + 7
        with Image.open(path) as source:
            ratio = source.height / max(source.width, 1)
        image_height = min(420, round((self.right - self.left - 90) * ratio))
        if self.page.y + image_height + estimated + 78 > self.content_bottom:
            self._new_page("content")
        self._picture(path, "界面  " + title, image_height)
        for label, value, _ in descriptions:
            prefix = label + "："
            lines = self._wrap(prefix + value, self.fonts["small"], self.right - self.left - 34)
            self._ensure(len(lines) * 30 + 7)
            page, top = self.page, self.page.y
            for index, line in enumerate(lines):
                self._draw_text(page, self.left + 18, page.y, line, self.fonts["small"],
                                ACCENT_DARK if index == 0 else INK)
                page.y += 30
            page.y += 7
            page.mark(top, page.y)

    def _draw_text(self, page: _Page, x: float, y: float, text: str, font,
                   color: str, bold: bool = False) -> None:
        page.draw.text((x, y), text, font=font, fill=color, stroke_width=1 if bold else 0,
                       stroke_fill=color)
        box = page.draw.textbbox((x, y), text, font=font, stroke_width=1 if bold else 0)
        page.mark(box[1], box[3])

    @staticmethod
    def _wrap(text: str, font, max_width: int) -> list:
        text = re.sub(r"\s+", " ", text.strip())
        if not text:
            return []
        probe = ImageDraw.Draw(Image.new("RGB", (1, 1)))
        lines, current = [], ""
        for character in text:
            candidate = current + character
            if current and probe.textlength(candidate, font=font) > max_width:
                lines.append(current.rstrip())
                current = character.lstrip()
            else:
                current = candidate
        if current:
            lines.append(current.rstrip())
        return lines

    @staticmethod
    def _asset(task_root: Path, relative: str) -> Optional[Path]:
        root = task_root.resolve()
        path = (root / relative).resolve()
        return path if root in path.parents and path.is_file() else None


class LibreOfficeManualRenderer:
    """Renders the actual exported DOCX instead of simulating Word pagination."""

    renderer_kind = "libreoffice_word"

    def __init__(self, office_renderer=None) -> None:
        # At 144 DPI an A4 page is exactly the existing 1191 × 1684 UI contract.
        self._office = office_renderer or LibreOfficeRenderer(dpi=144)

    @staticmethod
    def capability() -> dict:
        return LibreOfficeRenderer.capability()

    def render(self, document_path: Path, output_dir: Path, context: dict,
               sections: list, figures: list, screenshots: list,
               task_root: Path) -> CompanionRenderResult:
        try:
            rendered = self._office.render(document_path, output_dir)
        except SourceDocumentQaError as error:
            raise ManualQaError(str(error)) from error
        page_paths = tuple(rendered.page_paths)
        page_kinds = tuple(
            "cover" if index == 1 else "toc" if index == 2 else "content"
            for index in range(1, len(page_paths) + 1)
        )
        ratios = tuple(self._body_fill_ratio(path) for path in page_paths)
        last_page = len(page_paths)
        underfilled = tuple(
            index for index, (kind, ratio) in enumerate(
                zip(page_kinds, ratios), start=1
            ) if kind == "content" and (
                (index != last_page and ratio < 0.50)
                or (index == last_page and ratio < 0.25)
            )
        )
        return CompanionRenderResult(
            rendered.pdf_path, page_paths, page_kinds, ratios, underfilled
        )

    @staticmethod
    def _body_fill_ratio(path: Path) -> float:
        with Image.open(path) as source:
            image = source.convert("RGB")
        width, height = image.size
        # Exclude header and footer: they must not make an empty body page pass.
        body = image.crop((int(width * 0.07), int(height * 0.06),
                           int(width * 0.93), int(height * 0.94)))
        difference = ImageChops.difference(body, Image.new("RGB", body.size, "white"))
        bounds = difference.getbbox()
        if not bounds:
            return 0.0
        return round((bounds[3] - bounds[1]) / max(body.height, 1), 4)


class ManualDocxInspector:
    def __init__(self, font_asset: FontAsset = None) -> None:
        self.font_asset = font_asset or FontAsset.bundled_cjk()

    def inspect(self, document_path: Path, expected_sha256: str, sections: list,
                figures: list, screenshots: list,
                render: CompanionRenderResult, expectations: dict = None) -> ManualQaResult:
        expectations = expectations or {}
        checks = []
        actual_sha = hashlib.sha256(document_path.read_bytes()).hexdigest()
        checks.append(self._equal("artifact.sha256", expected_sha256, actual_sha))
        document = Document(document_path)
        section = document.sections[0]
        checks.append(self._equal("structure.a4_width", 7560310, section.page_width))
        checks.append(self._equal("structure.a4_height", 10692130, section.page_height))
        heading_count = sum(1 for paragraph in document.paragraphs
                            if paragraph.style.name == "Heading 1") - 1
        checks.append(self._equal("structure.chapter_headings", len(sections), heading_count))
        expected_tables = sum(1 for section_item in sections for block in section_item["blocks"]
                              if block.get("type") == "table"
                              and section_item.get("section_key") != "ui_operations")
        checks.append(self._equal("structure.tables", expected_tables, len(document.tables)))
        expected_images = len(figures) + len(screenshots)
        checks.append(self._minimum("structure.inline_images", expected_images,
                                    len(document.inline_shapes)))
        captions = [paragraph.text for paragraph in document.paragraphs
                    if paragraph.style.name == "Caption"]
        figure_numbers = [int(match.group(1)) for value in captions
                          if (match := re.match(r"^图 (\d+)\s+", value))]
        screenshot_numbers = figure_numbers[len(figures):]
        checks.append(self._equal(
            "structure.figure_caption_numbers",
            list(range(1, len(figures) + len(screenshots) + 1)), figure_numbers,
        ))
        checks.append(self._equal(
            "structure.screenshot_caption_numbers",
            list(range(len(figures) + 1, len(figures) + len(screenshots) + 1)),
            screenshot_numbers,
        ))
        with zipfile.ZipFile(document_path) as archive:
            names = archive.namelist()
            document_xml = archive.read("word/document.xml").decode("utf-8")
            font_table = archive.read("word/fontTable.xml").decode("utf-8")
            settings = archive.read("word/settings.xml").decode("utf-8")
            content_types = archive.read("[Content_Types].xml").decode("utf-8")
        checks.append(self._minimum("font.embedded_parts", 1,
                                    sum(1 for name in names if name.endswith(".odttf"))))
        checks.append(self._contains("font.embed_regular", font_table, "embedRegular"))
        checks.append(self._contains("font.embed_setting", settings, "embedTrueTypeFonts"))
        checks.append(self._contains("font.content_type", content_types, "obfuscatedFont"))
        checks.append(self._equal("structure.page_breaks", 2,
                                  document_xml.count('w:type="page"')))
        full_text = "".join(paragraph.text for paragraph in document.paragraphs)
        if expectations.get("document_kind") == "final_document":
            forbidden_final_markers = [marker for marker in (
                "章节导航（页码将在 Word 打开或导出时自动更新）",
                "图中内容以已审核的真实截图为准",
                "页面实际内容和操作范围以图中已审核的真实界面为准",
            ) if marker in full_text]
            checks.append(ManualQaCheck(
                "content.final_review_markers", not forbidden_final_markers, "blocker",
                [], forbidden_final_markers,
                "passed" if not forbidden_final_markers else "终稿仍包含仅供审阅阶段使用的提示语",
            ))
        font_summary = self.font_asset.validate(full_text)
        checks.append(self._equal("font.missing_codepoints", 0,
                                  font_summary["missing_codepoints"]))
        checks.append(self._minimum("render.page_count", 3, len(render.page_paths)))
        dimensions = []
        blank_pages = []
        for index, path in enumerate(render.page_paths, 1):
            with Image.open(path) as image:
                dimensions.append(image.size)
                extrema = image.convert("L").getextrema()
                if extrema == (255, 255):
                    blank_pages.append(index)
        checks.append(self._equal("render.a4_dimensions", True,
                                  all(size == A4_PIXELS for size in dimensions)))
        checks.append(self._equal("render.blank_pages", [], blank_pages))
        # Header and footer text means a visually empty body page is not a
        # pixel-perfect white image. Treat zero-density body pages as blockers
        # so an otherwise polished final document cannot pass with a numbered,
        # header-only sheet in the middle.
        body_blank_pages = [
            index for index, (kind, ratio) in enumerate(
                zip(render.page_kinds, render.fill_ratios), start=1
            ) if kind == "content" and ratio < 0.01
        ]
        checks.append(self._equal("render.body_blank_pages", [], body_blank_pages))
        if render.underfilled_pages:
            checks.append(ManualQaCheck(
                "render.page_density", False, "warning",
                "中间正文页填充率不低于 50%，末页不低于 25%",
                list(render.underfilled_pages),
                "存在内容较少的章节尾页或末页，建议复核分页，但不阻断交付",
            ))
        else:
            checks.append(ManualQaCheck(
                "render.page_density", True, "warning",
                "中间正文页填充率不低于 50%，末页不低于 25%", [], "passed",
            ))
        placeholder_hits = sorted(set(re.findall(
            r"待确认|待补充|TODO|TBD|后续迭代|建议后续", full_text, re.I
        )))
        checks.append(ManualQaCheck(
            "content.placeholders", not placeholder_hits, "blocker", [], placeholder_hits,
            "passed" if not placeholder_hits else "正文仍包含占位或路线图措辞，不能作为当前版本实现说明",
        ))
        epistemic_hits = sorted(set(re.findall(
            r"根据(?:项目|现有)?证据推断|据此推断|合理推断", full_text
        )))
        checks.append(ManualQaCheck(
            "content.epistemic_caveats", not epistemic_hits, "blocker", [], epistemic_hits,
            "passed" if not epistemic_hits else
            "正式材料中仍包含推断性措辞，必须改写为可由证据直接支持的事实",
        ))
        unverified_outcomes = unverified_outcome_hits(full_text)
        checks.append(ManualQaCheck(
            "content.unverified_outcomes", not unverified_outcomes, "blocker", [],
            unverified_outcomes,
            "passed" if not unverified_outcomes else
            "正文包含源码和项目材料无法直接证明的测试结果、验收或上线结论",
        ))
        source_name_candidates = set(
            match.group(0)
            for match in re.finditer(
                r"(?<![\w/])(?:[\w.-]+/)+[\w.-]+\.(?:py|java|kt|ts|tsx|js|jsx|vue|go|rs|cs|php)|"
                r"(?<![\w/])[A-Za-z][\w.-]*\.(?:py|java|kt|ts|tsx|js|jsx|vue|go|rs|cs|php)\b",
                full_text,
                re.I,
            )
        )
        # Product/framework names are legitimate prose, not source file names.
        framework_names = {"vue.js", "react.js", "node.js", "next.js", "nuxt.js"}
        source_file_mentions = sorted(
            value for value in source_name_candidates
            if value.lower() not in framework_names
        )
        checks.append(ManualQaCheck(
            "content.internal_source_names", not source_file_mentions, "warning", [],
            source_file_mentions,
            "passed" if not source_file_mentions else
            "正式叙述仍暴露内部源码文件名；应改写为中文业务角色并仅在证据元数据中保留路径",
        ))
        unsupported_blocks = []
        substantive_block_count = 0
        evidence_bound_block_count = 0
        for item in sections:
            for block_index, block in enumerate(item.get("blocks", []), start=1):
                if block.get("type") in {"figure_request", "subheading"}:
                    continue
                substantive_block_count += 1
                refs = [str(ref).strip() for ref in block.get("evidence_refs", [])
                        if str(ref).strip()]
                if refs:
                    evidence_bound_block_count += 1
                else:
                    unsupported_blocks.append({
                        "section_key": item.get("section_key"),
                        "block_index": block_index,
                        "block_type": block.get("type"),
                    })
        checks.append(ManualQaCheck(
            "content.evidence_coverage", not unsupported_blocks, "blocker",
            "每个正文、列表和表格块至少绑定 1 条真实证据",
            {
                "substantive_blocks": substantive_block_count,
                "evidence_bound_blocks": evidence_bound_block_count,
                "unsupported_blocks": unsupported_blocks,
            },
            "passed" if not unsupported_blocks else "存在未绑定证据的事实性内容块",
        ))
        inference_blocks = [
            {"section_key": item.get("section_key"), "block_index": index}
            for item in sections
            for index, block in enumerate(item.get("blocks", []), start=1)
            if block.get("inference")
        ]
        checks.append(ManualQaCheck(
            "content.inference_claims", not inference_blocks, "blocker", [],
            inference_blocks,
            "passed" if not inference_blocks else
            "正文仍包含 AI 推断，必须改写为源码可直接证明的实现事实",
        ))
        review_sections = [
            item.get("section_key") for item in sections
            if item.get("status") == "needs_review"
        ]
        checks.append(ManualQaCheck(
            "content.section_review_status", not review_sections, "blocker",
            "所有章节均已通过证据规范化或人工确认", review_sections,
            "passed" if not review_sections else "部分章节仍需复核，不能标记为正式通过版",
        ))
        section_quality = []
        for item in sections:
            # Chapter 7 is delivered from the reviewed screenshot snapshot, not
            # from generic AI prose volume. Screenshot authenticity, six-part
            # interpretation, order and sensitive-data review have dedicated
            # blocker checks below, so applying the 950-character chapter rule
            # here double-counts the wrong evidence and penalizes concise pages.
            if item.get("section_key") == "ui_operations" and screenshots:
                continue
            substantive = [block for block in item["blocks"]
                           if block.get("type") not in {"figure_request", "subheading"}]
            paragraphs = [block for block in substantive
                          if block.get("type") == "paragraph"]
            paragraph_characters = sum(len(str(block.get("text", "")))
                                       for block in paragraphs)
            content = "".join(
                str(block.get("text", "")) + str(block.get("lead", "")) +
                "".join(str(value) for value in block.get("items", [])) +
                "".join(str(cell) for row in block.get("rows", []) for cell in row)
                for block in substantive
            )
            minimum = 750 if item.get("section_key") == "introduction" else 950
            if (len(content) < minimum or len(substantive) < 4 or
                    len(paragraphs) < 2 or paragraph_characters < 320):
                section_quality.append({
                    "section_key": item.get("section_key"),
                    "characters": len(content), "blocks": len(substantive),
                    "paragraphs": len(paragraphs),
                    "paragraph_characters": paragraph_characters,
                })
        checks.append(ManualQaCheck(
            "content.section_depth", not section_quality, "blocker",
            "每章至少 4 个实质块和 2 个正文段落；引言不少于 750 字，其他章节不少于 950 字",
            section_quality,
            "passed" if not section_quality else "部分章节仍像提纲，正文深度不足",
        ))
        expected_section_keys = set(expectations.get("section_keys") or [])
        actual_section_keys = {item.get("section_key") for item in sections}
        missing_sections = sorted(expected_section_keys - actual_section_keys)
        checks.append(ManualQaCheck(
            "content.required_sections", not missing_sections, "blocker",
            sorted(expected_section_keys), sorted(actual_section_keys),
            "passed" if not missing_sections else
            "项目证据表明这些章节适用，但正文生成失败或缺失：{0}".format(
                "、".join(missing_sections)
            ),
        ))
        expected_figure_sections = set(expectations.get("figure_sections") or [])
        actual_figure_sections = {item.get("section_key") for item in figures
                                  if item.get("png_relative_path")}
        missing_figures = sorted(expected_figure_sections - actual_figure_sections)
        checks.append(ManualQaCheck(
            "content.figure_coverage", not missing_figures, "blocker",
            sorted(expected_figure_sections), sorted(actual_figure_sections),
            "passed" if not missing_figures else
            "必要章节图表尚未成功生成：{0}".format("、".join(missing_figures)),
        ))
        ui_text = "".join(
            json.dumps(item.get("blocks", []), ensure_ascii=False)
            for item in sections if item.get("section_key") == "ui_operations"
        )
        ui_applicable = bool(expectations.get("ui_applicable")) or bool(re.search(
            r"界面|页面|按钮|前端|vue|react|客户端", ui_text, re.I
        ))
        ui_mode = expectations.get("ui_evidence_mode") or "unknown"
        checks.append(ManualQaCheck(
            "content.ui_screenshot", not ui_applicable or bool(screenshots), "blocker",
            "识别到真实界面时至少包含 1 张已说明的截图",
            {"count": len(screenshots), "mode": ui_mode},
            "passed" if not ui_applicable or screenshots else
            ({"waiting_for_screenshots": "用户界面章节正在等待用户提供并确认真实截图证据",
              "not_applicable": "用户已确认项目不适用截图",
              "source_inferred": "用户明确选择了源码推断版，正式质量仍需人工确认"}
             .get(ui_mode, "用户界面章节尚未获得可采用的真实截图证据")),
        ))
        checks.append(ManualQaCheck(
            "content.screenshot_interpretations_reviewed",
            all(item.get("interpretation_revision_id") for item in screenshots), "blocker",
            "所有采用截图均绑定明确的已审核解读版本",
            [item.get("interpretation_revision_id") for item in screenshots],
            "passed" if all(item.get("interpretation_revision_id") for item in screenshots)
            else "存在未绑定已审核解读版本的采用截图",
        ))
        screenshot_refs = ["screenshot:{0}:v{1}".format(
            item.get("asset_id"), item.get("interpretation_version")
        ) for item in screenshots if item.get("asset_id") and item.get("interpretation_version")]
        ui_blocks = [block for section in sections if section.get("section_key") == "ui_operations"
                     for block in section.get("blocks", [])]
        first_seen_refs, invalid_ui_blocks = [], []
        valid_screenshot_refs = set(screenshot_refs)
        for index, block in enumerate(ui_blocks):
            refs = [ref for ref in block.get("evidence_refs", [])
                    if str(ref).startswith("screenshot:")]
            if (screenshot_refs and block.get("type") not in {"subheading", "figure_request"} and
                    (not refs or any(ref not in valid_screenshot_refs for ref in refs))):
                invalid_ui_blocks.append(index)
            for ref in refs:
                if ref in valid_screenshot_refs and ref not in first_seen_refs:
                    first_seen_refs.append(ref)
        checks.append(ManualQaCheck(
            "content.ui_screenshot_evidence_refs", not screenshot_refs or not invalid_ui_blocks,
            "blocker", "第 7 章事实性内容只引用当前采用截图解读版本",
            {"invalid_block_indexes": invalid_ui_blocks, "valid_refs": screenshot_refs},
            "passed" if not screenshot_refs or not invalid_ui_blocks else
            "第 7 章存在未绑定当前采用截图解读版本的事实性内容",
        ))
        checks.append(ManualQaCheck(
            "content.ui_group_order", not screenshot_refs or first_seen_refs == screenshot_refs,
            "blocker", screenshot_refs, first_seen_refs,
            "passed" if not screenshot_refs or first_seen_refs == screenshot_refs else
            "第 7 章的截图引用顺序与人工确认的页面组和截图顺序不一致",
        ))
        checks.append(ManualQaCheck(
            "content.ui_section_current",
            not screenshots or bool(expectations.get("ui_section_current")), "blocker",
            "用户界面章节基于当前采用截图版本生成",
            expectations.get("ui_evidence_mode"),
            "passed" if not screenshots or expectations.get("ui_section_current")
            else "截图、解读或项目概要已变化，用户界面章节需要更新",
        ))
        checks.append(ManualQaCheck(
            "content.screenshot_sensitive_review",
            all(item.get("sensitive_status", "confirmed_safe") == "confirmed_safe"
                for item in screenshots), "blocker", "采用截图的敏感信息已确认",
            [item.get("sensitive_status") for item in screenshots],
            "passed" if all(item.get("sensitive_status", "confirmed_safe") == "confirmed_safe"
                            for item in screenshots)
            else "存在尚未确认敏感信息或已标记含敏感信息的截图",
        ))
        screenshot_reference_numbers = [number for number in range(
            len(figures) + 1, len(figures) + len(screenshots) + 1
        ) if re.search(r"图\s*{0}(?!\d)".format(number), full_text)]
        checks.append(self._equal(
            "content.screenshot_figure_references",
            list(range(len(figures) + 1, len(figures) + len(screenshots) + 1)),
            screenshot_reference_numbers,
        ))
        screenshot_signatures = {}
        for item in screenshots:
            signature = (
                str(item.get("title", "")).strip(),
                json.dumps(item.get("description", {}), ensure_ascii=False, sort_keys=True),
            )
            screenshot_signatures.setdefault(signature, []).append(
                item.get("screenshot_key") or item.get("image_relative_path")
            )
        repeated_screenshots = [values for values in screenshot_signatures.values()
                                if len(values) > 1]
        checks.append(ManualQaCheck(
            "content.screenshot_distinctness", not repeated_screenshots, "blocker",
            "每张截图具有不同页面标题和针对该页面的六维说明",
            repeated_screenshots,
            "passed" if not repeated_screenshots else
            "存在标题和说明完全相同的截图；不能用同一空页面重复充当多个操作界面",
        ))
        passed = all(check.passed for check in checks if check.severity == "blocker")
        warning_count = sum(1 for check in checks
                            if check.severity == "warning" and not check.passed)
        failed = [check.key for check in checks
                  if check.severity == "blocker" and not check.passed]
        summary = {
            "passed": passed,
            "policy_version": QA_POLICY_VERSION,
            "renderer_kind": RENDERER_KIND,
            "renderer_disclosure": "运行时页面由同源确定性渲染器生成。",
            "word_render_baseline": "companion_rendered",
            "check_count": len(checks),
            "failed_check_count": len(failed),
            "failed_checks": failed,
            "warning_count": warning_count,
            "rendered_pages": len(render.page_paths),
            "underfilled_pages": list(render.underfilled_pages),
            "page_fill_ratios": list(render.fill_ratios),
            "page_dimensions": list(A4_PIXELS),
            "document_sha256": actual_sha,
            "font": font_summary,
        }
        return ManualQaResult(passed, tuple(checks), summary, render)

    @staticmethod
    def _equal(key: str, expected: object, actual: object) -> ManualQaCheck:
        passed = expected == actual
        return ManualQaCheck(key, passed, "blocker", expected, actual,
                             "passed" if passed else "expected {0!r}, got {1!r}".format(expected, actual))

    @staticmethod
    def _minimum(key: str, minimum: int, actual: int) -> ManualQaCheck:
        passed = actual >= minimum
        return ManualQaCheck(key, passed, "blocker", ">= {0}".format(minimum), actual,
                             "passed" if passed else "expected >= {0}, got {1}".format(minimum, actual))

    @staticmethod
    def _contains(key: str, text: str, needle: str) -> ManualQaCheck:
        return ManualDocxInspector._equal(key, True, needle in text)


class ManualQaService:
    def __init__(self, database: Database, data_root: Path, *, documents=None,
                 renderer=None, inspector=None) -> None:
        self._database = database
        self._data_root = data_root.expanduser().resolve()
        self._documents = documents or ManualDocumentService(database, data_root)
        self._renderer = renderer or LibreOfficeManualRenderer()
        self._inspector = inspector or ManualDocxInspector()

    def _coverage_expectations(self, job_id: str, sections: list) -> dict:
        section_keys = {item.get("section_key") for item in sections}
        has_ui = "ui_operations" in section_keys
        has_research = False
        with self._database.connect() as connection:
            row = connection.execute(
                """SELECT project_profile_json,notes_json FROM manual_research_artifacts
                WHERE job_id=? ORDER BY version DESC LIMIT 1""", (job_id,),
            ).fetchone()
            ui_node = connection.execute(
                """SELECT status,output_json FROM manual_execution_nodes
                WHERE job_id=? AND node_key='section:ui_operations'""", (job_id,),
            ).fetchone()
            assessment = connection.execute(
                """SELECT status FROM manual_capture_assessments
                WHERE job_id=? ORDER BY version DESC LIMIT 1""", (job_id,),
            ).fetchone()
        if row is not None:
            has_research = True
            try:
                profile = json.loads(row["project_profile_json"] or "{}")
                notes = json.loads(row["notes_json"] or "{}")
            except json.JSONDecodeError:
                profile, notes = {}, {}
            guidance_keys = {
                item.get("section_key") for item in notes.get("section_guidance", [])
                if isinstance(item, dict)
            }
            profile_text = json.dumps(profile, ensure_ascii=False).lower()
            has_ui = "ui_operations" in guidance_keys or any(
                marker in profile_text for marker in
                ("vue", "react", "svelte", "tauri", "electron", "frontend", "前端", "界面")
            )
        if has_research:
            expected_sections = {
                "introduction", "architecture", "modules", "data_interfaces",
                "runtime", "security_reliability", "testing_summary",
            }
            if has_ui:
                expected_sections.add("ui_operations")
            required_figure_sections = {
                "architecture", "modules", "data_interfaces", "runtime",
            }
        else:
            expected_sections = section_keys
            required_figure_sections = {
                item.get("section_key") for item in sections
                if any(block.get("type") == "figure_request"
                       for block in item.get("blocks", []))
            }
        ui_output = json.loads(ui_node["output_json"] or "{}") if ui_node else {}
        explicit_mode = ui_output.get("evidence_mode")
        return {
            "section_keys": sorted(expected_sections),
            "figure_sections": sorted(required_figure_sections),
            "ui_applicable": has_ui,
            "ui_evidence_mode": (
                explicit_mode if explicit_mode in {"source_inferred", "not_applicable"}
                else "not_applicable" if assessment and assessment["status"] == "not_applicable"
                else "waiting_for_screenshots" if ui_node and ui_node["status"] in {
                    "waiting_for_screenshots", "waiting_for_review", "outdated"
                } else "screenshot_driven" if ui_node and ui_node["status"] == "completed"
                else "unknown"
            ),
            "ui_section_current": bool(not ui_node or ui_node["status"] == "completed"),
        }

    @staticmethod
    def _rendered_toc_pages(pdf_path: Path, sections: list) -> dict:
        """Read chapter heading pages from the rendered PDF when Poppler is available."""
        candidates = [shutil.which("pdftotext")]
        _, pdftoppm = LibreOfficeRenderer._resolve_tools()
        if pdftoppm:
            executable = "pdftotext.exe" if Path(pdftoppm).suffix.lower() == ".exe" \
                else "pdftotext"
            candidates.append(str(Path(pdftoppm).with_name(executable)))
        candidates.extend(["/opt/homebrew/bin/pdftotext", "/usr/local/bin/pdftotext"])
        pdftotext = next((item for item in candidates if item and Path(item).is_file()), None)
        if not pdftotext:
            return {}
        completed = subprocess.run(
            [pdftotext, "-layout", str(pdf_path), "-"], capture_output=True,
            timeout=60, check=False,
        )
        if completed.returncode != 0:
            return {}
        pages = completed.stdout.decode("utf-8", errors="replace").split("\f")
        result = {}
        for index, section in enumerate(sections, 1):
            heading = re.compile(
                r"(?m)^\s*{0}\s+{1}\s*$".format(
                    index, re.escape(str(section.get("title", "")).strip())
                )
            )
            for page_number, page in enumerate(pages, 1):
                if page_number > 2 and heading.search(page):
                    result[index] = page_number
                    break
        return result if len(result) == len(sections) else {}

    def execute(self, job_id: str, version: Optional[int] = None) -> dict:
        self._database.initialize()
        preview = self._documents.preview(job_id, version or self._documents.get(job_id)["version"])
        document = preview["document"]
        if document.get("document_kind") == "review_checkpoint":
            raise ManualQaError("阶段审阅稿仅用于正文预览与落盘；请等待正式候选稿后再执行逐页质检")
        if document["integrity"]["status"] != "verified":
            raise ManualQaError("正式说明书文件缺失或完整性校验失败")
        step_id = self._start_step(job_id)
        task_root = self._data_root / "tasks" / document["task_id"]
        document_path = task_root / document["docx_relative_path"]
        qa_version = self._next_version(document["id"])
        base_relative = Path("qa") / "manual" / job_id / "doc-v{0}".format(document["version"])
        render_relative = base_relative / "render-v{0}".format(qa_version)
        report_relative = base_relative / "qa-v{0}.json".format(qa_version)
        render_path = task_root / render_relative
        report_path = task_root / report_relative
        context = {
            "project_name": document["project_name"],
            "project_version": document["project_version"],
        }
        try:
            runtime_renderer = getattr(self._renderer, "renderer_kind", RENDERER_KIND)
            if runtime_renderer == "libreoffice_word":
                probe_path = render_path / "toc-probe"
                probe = self._renderer.render(
                    document_path, probe_path, context, preview["sections"], preview["figures"],
                    preview["screenshots"], task_root,
                )
                toc_pages = self._rendered_toc_pages(probe.pdf_path, preview["sections"])
                if toc_pages:
                    document = self._documents.persist_toc_page_results(
                        job_id, document["version"], toc_pages
                    )
                render = self._renderer.render(
                    document_path, render_path, context, preview["sections"], preview["figures"],
                    preview["screenshots"], task_root,
                )
                shutil.rmtree(probe_path, ignore_errors=True)
            else:
                render = self._renderer.render(
                    document_path, render_path, context, preview["sections"], preview["figures"],
                    preview["screenshots"], task_root,
                )
            expectations = self._coverage_expectations(job_id, preview["sections"])
            expectations["document_kind"] = document.get("document_kind")
            result = self._inspector.inspect(
                document_path, document["sha256"], preview["sections"],
                preview["figures"], preview["screenshots"], render, expectations,
            )
            result.summary["renderer_kind"] = runtime_renderer
            if runtime_renderer == "libreoffice_word":
                result.summary["renderer_disclosure"] = (
                    "运行时页面由 LibreOffice 直接渲染当前导出的 DOCX。"
                )
                result.summary["word_render_baseline"] = "current_document_rendered"
            payload = {
                "schema_version": 1,
                "job_id": job_id,
                "document_artifact_id": document["id"],
                "document_version": document["version"],
                "qa_version": qa_version,
                "summary": result.summary,
                "checks": [check.__dict__ for check in result.checks],
            }
            report_path.parent.mkdir(parents=True, exist_ok=True)
            report_path.write_text(
                json.dumps(payload, ensure_ascii=False, sort_keys=True, indent=2) + "\n",
                encoding="utf-8",
            )
        except Exception as error:
            self._finish_failure(job_id, step_id, error)
            if isinstance(error, (ManualQaError, ManualDocumentError)):
                raise
            raise ManualQaError("说明书质量检查失败，已保留 DOCX 和既有阶段结果") from error
        now = utc_now()
        status = "qa_passed" if result.passed else "qa_failed"
        step_status = "completed" if result.passed and not result.summary["warning_count"] \
            else "completed_with_warnings"
        job_status = "completed" if step_status == "completed" else "completed_with_warnings"
        merged_qa = dict(document["qa"])
        merged_qa["quality"] = result.summary
        merged_qa["warning_count"] = (
            int(merged_qa.get("warning_count", 0)) + result.summary["warning_count"]
        )
        preview_pdf_relative = (render_relative / "preview.pdf").as_posix()
        qa_run_id = str(uuid4())
        with self._database.connect() as connection:
            connection.execute(
                """INSERT INTO manual_document_qa_runs(id,document_artifact_id,job_id,
                qa_version,policy_version,renderer_kind,passed,checks_json,summary_json,
                report_relative_path,render_relative_path,preview_pdf_relative_path,created_at)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                (qa_run_id, document["id"], job_id, qa_version, QA_POLICY_VERSION,
                 runtime_renderer, 1 if result.passed else 0,
                 json.dumps([check.__dict__ for check in result.checks], ensure_ascii=False,
                            separators=(",", ":")),
                 json.dumps(result.summary, ensure_ascii=False, separators=(",", ":")),
                 report_relative.as_posix(), render_relative.as_posix(),
                 preview_pdf_relative, now),
            )
            connection.execute(
                """UPDATE manual_document_artifacts SET status=?,preview_pdf_relative_path=?,
                qa_json=? WHERE id=?""",
                (status, preview_pdf_relative,
                 json.dumps(merged_qa, ensure_ascii=False, separators=(",", ":")),
                 document["id"]),
            )
            connection.execute(
                """UPDATE manual_generation_steps SET status=?,summary_json=?,finished_at=?,
                safe_error_message=NULL WHERE id=?""",
                (step_status, json.dumps(result.summary, ensure_ascii=False,
                                         separators=(",", ":")), now, step_id),
            )
            connection.execute(
                """UPDATE manual_generation_jobs SET status=?,current_step='render_qa',
                progress_json=?,finished_at=?,updated_at=?,safe_error_message=NULL WHERE id=?""",
                (job_status, json.dumps({"completed": 6, "total": 6, "percent": 100},
                                        separators=(",", ":")), now, now, job_id),
            )
        return {"document": self._documents.get(job_id, document["version"]),
                "qa_run": self.get(job_id, document["version"])}

    def get(self, job_id: str, version: int) -> dict:
        document = self._documents.get(job_id, version)
        with self._database.connect() as connection:
            row = connection.execute(
                """SELECT q.* FROM manual_document_qa_runs q
                WHERE q.document_artifact_id=? ORDER BY q.qa_version DESC LIMIT 1""",
                (document["id"],),
            ).fetchone()
            decisions = [] if row is None else connection.execute(
                """SELECT check_key,action,reason,created_at
                FROM manual_qa_decisions WHERE qa_run_id=? ORDER BY created_at""",
                (row["id"],),
            ).fetchall()
        if row is None:
            raise ManualQaError("该说明书版本尚未执行质量检查")
        return {
            "id": row["id"], "job_id": row["job_id"],
            "document_artifact_id": row["document_artifact_id"],
            "document_version": version, "qa_version": row["qa_version"],
            "policy_version": row["policy_version"], "renderer_kind": row["renderer_kind"],
            "passed": bool(row["passed"]), "checks": json.loads(row["checks_json"]),
            "summary": json.loads(row["summary_json"]),
            "page_count": json.loads(row["summary_json"])["rendered_pages"],
            "decisions": [dict(item) for item in decisions],
            "created_at": row["created_at"],
        }

    def defer_check(self, job_id: str, version: int, check_key: str, reason: str) -> dict:
        """Waive one eligible failure while preserving the original QA facts."""
        self._database.initialize()
        reason = reason.strip()
        if len(reason) < 4:
            raise ManualQaError("请说明本轮忽略该问题的原因")
        qa = self.get(job_id, version)
        check = next((item for item in qa["checks"] if item["key"] == check_key), None)
        if check is None:
            raise ManualQaError("质量检查项不存在")
        if check["passed"]:
            raise ManualQaError("该质量检查项已经通过，无需忽略")
        waivable = {
            "render.page_density",
            "content.section_depth",
            "content.figure_coverage",
        }
        if check_key not in waivable:
            raise ManualQaError("该项涉及文档真实性、完整性或安全性，不能通过忽略绕过")
        now = utc_now()
        with self._database.connect() as connection:
            connection.execute(
                """INSERT INTO manual_qa_decisions(
                id,qa_run_id,document_artifact_id,job_id,check_key,action,reason,created_at)
                VALUES (?,?,?,?,?,'deferred',?,?)
                ON CONFLICT(qa_run_id,check_key) DO UPDATE SET
                action='deferred',reason=excluded.reason,created_at=excluded.created_at""",
                (str(uuid4()), qa["id"], qa["document_artifact_id"], job_id,
                 check_key, reason, now),
            )
            decisions = connection.execute(
                """SELECT check_key FROM manual_qa_decisions
                WHERE qa_run_id=? AND action='deferred'""", (qa["id"],),
            ).fetchall()
            waived = sorted({row["check_key"] for row in decisions})
            raw_failed = sorted(item["key"] for item in qa["checks"]
                                if item["severity"] == "blocker" and not item["passed"])
            effective_failed = [key for key in raw_failed if key not in waived]
            effective_passed = not effective_failed
            summary = dict(qa["summary"])
            summary.update({
                "passed": effective_passed,
                "raw_failed_checks": raw_failed,
                "waived_checks": waived,
                "waived_check_count": len(waived),
                "failed_checks": effective_failed,
                "failed_check_count": len(effective_failed),
            })
            document = connection.execute(
                """SELECT qa_json FROM manual_document_artifacts WHERE id=?""",
                (qa["document_artifact_id"],),
            ).fetchone()
            document_qa = json.loads(document["qa_json"] or "{}")
            document_qa["quality"] = summary
            connection.execute(
                """UPDATE manual_document_qa_runs SET passed=?,summary_json=? WHERE id=?""",
                (1 if effective_passed else 0,
                 json.dumps(summary, ensure_ascii=False, separators=(",", ":")), qa["id"]),
            )
            connection.execute(
                """UPDATE manual_document_artifacts SET status=?,qa_json=? WHERE id=?""",
                ("qa_passed" if effective_passed else "qa_failed",
                 json.dumps(document_qa, ensure_ascii=False, separators=(",", ":")),
                 qa["document_artifact_id"]),
            )
            connection.execute(
                """UPDATE manual_generation_jobs SET status=?,updated_at=? WHERE id=?""",
                ("completed_with_warnings" if waived or not effective_passed else "completed",
                 now, job_id),
            )
        return self.get(job_id, version)

    def read_page(self, job_id: str, version: int, page_number: int) -> bytes:
        item, document = self.get(job_id, version), self._documents.get(job_id, version)
        if page_number < 1 or page_number > item["page_count"]:
            raise ManualQaError("预览页码超出范围")
        with self._database.connect() as connection:
            row = connection.execute(
                """SELECT render_relative_path FROM manual_document_qa_runs
                WHERE id=?""", (item["id"],),
            ).fetchone()
        render_path = (self._data_root / "tasks" / document["task_id"] /
                       Path(row["render_relative_path"])).resolve()
        page_path = next((candidate for candidate in render_path.glob("page-*.png")
                          if candidate.stem.split("-")[-1].isdigit() and
                          int(candidate.stem.split("-")[-1]) == page_number), None)
        if page_path is None:
            raise ManualQaError("质量检查预览页不存在")
        return self._safe_path(
            document["task_id"], page_path.relative_to(
                (self._data_root / "tasks" / document["task_id"]).resolve()
            )
        ).read_bytes()

    def read_pdf(self, job_id: str, version: int) -> bytes:
        item, document = self.get(job_id, version), self._documents.get(job_id, version)
        with self._database.connect() as connection:
            row = connection.execute(
                """SELECT preview_pdf_relative_path FROM manual_document_qa_runs
                WHERE id=?""", (item["id"],),
            ).fetchone()
        return self._safe_path(document["task_id"], Path(row["preview_pdf_relative_path"])).read_bytes()

    def _safe_path(self, task_id: str, relative: Path) -> Path:
        root = (self._data_root / "tasks" / task_id).resolve()
        path = (root / relative).resolve()
        if root not in path.parents or not path.is_file():
            raise ManualQaError("质量检查产物不存在或路径无效")
        return path

    def _next_version(self, document_id: str) -> int:
        with self._database.connect() as connection:
            return connection.execute(
                """SELECT COALESCE(MAX(qa_version),0)+1 value
                FROM manual_document_qa_runs WHERE document_artifact_id=?""",
                (document_id,),
            ).fetchone()["value"]

    def _start_step(self, job_id: str) -> str:
        now = utc_now()
        with self._database.connect() as connection:
            row = connection.execute(
                """SELECT id,status,attempt FROM manual_generation_steps WHERE job_id=?
                AND step_key='render_qa' ORDER BY attempt DESC LIMIT 1""", (job_id,),
            ).fetchone()
            if row is None:
                raise ManualQaError("说明书任务缺少质量检查阶段")
            if row["status"] == "running":
                raise ManualQaError("说明书正在执行质量检查，请勿重复提交")
            if row["status"] in {"completed", "completed_with_warnings"}:
                step_id = str(uuid4())
                connection.execute(
                    """INSERT INTO manual_generation_steps(id,job_id,step_key,status,attempt,
                    summary_json,started_at) VALUES (?,?,'render_qa','running',?,'{}',?)""",
                    (step_id, job_id, row["attempt"] + 1, now),
                )
            else:
                step_id = row["id"]
                connection.execute(
                    """UPDATE manual_generation_steps SET status='running',started_at=?,
                    finished_at=NULL,safe_error_message=NULL WHERE id=?""", (now, step_id),
                )
            connection.execute(
                """UPDATE manual_generation_jobs SET status='running',current_step='render_qa',
                updated_at=?,safe_error_message=NULL WHERE id=?""", (now, job_id),
            )
        return step_id

    def _finish_failure(self, job_id: str, step_id: str, error: Exception) -> None:
        now = utc_now()
        message = "说明书质量检查失败：{0}".format(type(error).__name__)
        with self._database.connect() as connection:
            connection.execute(
                """UPDATE manual_generation_steps SET status='failed',finished_at=?,
                safe_error_message=? WHERE id=?""", (now, message, step_id),
            )
            connection.execute(
                """UPDATE manual_generation_jobs SET status='failed',current_step='render_qa',
                updated_at=?,safe_error_message=? WHERE id=?""", (now, message, job_id),
            )

import hashlib
import json
import os
import re
import tempfile
import zipfile
from pathlib import Path
from typing import Optional
from uuid import uuid4

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

from .service import utc_now
from .storage import Database


GENERATOR_VERSION = "formal-manual-docx-v12"
BODY_LINE_SPACING = 1.18
BODY_SPACE_AFTER_PT = 4.0
FONT_NAME = "Noto Sans CJK SC"
FONT_FILE = Path(__file__).resolve().parent / "assets" / "fonts" / "noto-cjk" / "NotoSansCJKsc-Regular.otf"
INK = "172331"
# Registration materials use a restrained monochrome hierarchy. Earlier blue
# headings and labels made the screenshot chapter look visually disconnected.
ACCENT = INK
ACCENT_DARK = INK
MUTED = INK
TABLE_FILL = "E8EEF5"
TABLE_WIDTH_DXA = 9468
TABLE_INDENT_DXA = 120


class ManualDocumentError(ValueError):
    pass


class ManualDocumentService:
    """Assembles versioned, native Word manuals from formal pipeline artifacts."""

    def __init__(self, database: Database, data_root: Path) -> None:
        self._database = database
        self._data_root = data_root.expanduser().resolve()

    def assemble(self, job_id: str) -> dict:
        """Assemble the current review candidate and advance the formal pipeline."""
        return self._assemble(job_id, document_kind="formal_candidate", advance_pipeline=True)

    def finalize(self, job_id: str, source_version: int) -> dict:
        """Create a clean final document only after an explicit human decision."""
        source = self.get(job_id, source_version)
        if source["document_kind"] not in {"formal_candidate", "final_document"}:
            raise ManualDocumentError("只有已审阅候选稿或已有终稿可以重新生成人工终稿")
        if source["integrity"]["status"] != "verified":
            raise ManualDocumentError("审阅稿文件缺失或完整性校验失败")
        if source["freshness"]["status"] != "current":
            raise ManualDocumentError("正文、图表或截图已有更新，请先重新装配审阅稿")
        if source["quality"]["status"] in {"not_checked", "outdated"}:
            raise ManualDocumentError("请先完成当前审阅稿的逐页检查，再由人工决定是否生成终稿")
        return self._assemble(
            job_id, document_kind="final_document", advance_pipeline=False,
            source_document_version=source_version,
        )

    def assemble_checkpoint(self, job_id: str) -> dict:
        """Persist a prose-only review checkpoint without advancing formal stages.

        Checkpoints intentionally snapshot only completed section content.  Figures and
        screenshots continue on independent nodes and can be included by a later formal
        candidate without mutating this immediately reviewable artifact.
        """
        return self._assemble(
            job_id, document_kind="review_checkpoint", advance_pipeline=False,
            include_assets=False,
        )

    def _assemble(self, job_id: str, *, document_kind: str,
                  advance_pipeline: bool, include_assets: bool = True,
                  source_document_version: Optional[int] = None) -> dict:
        self._database.initialize()
        context = self._context(job_id)
        sections = self._sections(job_id)
        if len(sections) < 3:
            raise ManualDocumentError("结构化正文不足，至少需要三个已生成章节")
        figures = self._figures(job_id) if include_assets else []
        screenshots = self._screenshots(job_id) if include_assets else []
        step_id = self._start_step(job_id) if advance_pipeline else None
        version = self._next_version(job_id)
        # Document versions restart at v1 for every generation job. Keep each job in
        # its own directory so a later job can never overwrite an earlier v1 file.
        stem = ({"formal_candidate": "software-manual-review",
                 "final_document": "software-manual-final"}.get(
                     document_kind, "review-checkpoint"))
        relative = (Path("artifacts") / "manual" / "jobs" /
                    "job-v{0}".format(context["job_version"]) /
                    "{0}.v{1}.docx".format(stem, version))
        task_root = self._data_root / "tasks" / context["task_id"]
        artifact = task_root / relative
        artifact.parent.mkdir(parents=True, exist_ok=True)
        descriptor, temporary_name = tempfile.mkstemp(
            prefix="manual-docx-", suffix=".docx", dir=str(artifact.parent)
        )
        os.close(descriptor)
        temporary = Path(temporary_name)
        try:
            build_context = dict(context)
            build_context["document_kind"] = document_kind
            build_context["source_document_version"] = source_document_version
            summary = FormalManualBuilder().build(
                temporary, build_context, sections, figures, screenshots, task_root
            )
            summary["document_kind"] = document_kind
            if source_document_version is not None:
                summary["source_document_version"] = source_document_version
            digest = hashlib.sha256(temporary.read_bytes()).hexdigest()
            os.replace(str(temporary), str(artifact))
        except Exception as error:
            temporary.unlink(missing_ok=True)
            if step_id is not None:
                self._finish_failure(job_id, step_id, error)
            if isinstance(error, ManualDocumentError):
                raise
            raise ManualDocumentError("正式说明书 DOCX 装配失败，已有阶段结果已保留") from error
        now = utc_now()
        with self._database.connect() as connection:
            connection.execute(
                """INSERT INTO manual_document_artifacts(id, job_id, version, status,
                docx_relative_path, preview_pdf_relative_path, qa_json, sha256, created_at)
                VALUES (?, ?, ?, 'assembled', ?, NULL, ?, ?, ?)""",
                (str(uuid4()), job_id, version, relative.as_posix(),
                 json.dumps(summary, ensure_ascii=False, separators=(",", ":")), digest, now),
            )
            if step_id is not None:
                connection.execute(
                    """UPDATE manual_generation_steps SET status='completed', summary_json=?,
                    finished_at=?, safe_error_message=NULL WHERE id=?""",
                    (json.dumps(summary, ensure_ascii=False, separators=(",", ":")),
                     now, step_id),
                )
                connection.execute(
                    """UPDATE manual_generation_jobs SET status='running', current_step='render_qa',
                    progress_json=?, updated_at=?, safe_error_message=NULL WHERE id=?""",
                    (json.dumps({"completed": 5, "total": 6, "percent": 83},
                                separators=(",", ":")), now, job_id),
                )
        return self.get(job_id, version)

    def list(self, job_id: str) -> list:
        self._database.initialize()
        context = self._context(job_id)
        with self._database.connect() as connection:
            rows = connection.execute(
                """SELECT * FROM manual_document_artifacts WHERE job_id=?
                ORDER BY version DESC""", (job_id,),
            ).fetchall()
        return [self._document_dict(context, row) for row in rows]

    def get(self, job_id: str, version: Optional[int] = None) -> dict:
        self._database.initialize()
        context = self._context(job_id)
        with self._database.connect() as connection:
            if version is None:
                row = connection.execute(
                    """SELECT * FROM manual_document_artifacts WHERE job_id=?
                    ORDER BY version DESC LIMIT 1""", (job_id,),
                ).fetchone()
            else:
                row = connection.execute(
                    """SELECT * FROM manual_document_artifacts WHERE job_id=? AND version=?""",
                    (job_id, version),
                ).fetchone()
        if row is None:
            raise ManualDocumentError("正式说明书版本不存在")
        return self._document_dict(context, row)

    def read(self, job_id: str, version: int) -> bytes:
        item = self.get(job_id, version)
        if item["integrity"]["status"] != "verified":
            raise ManualDocumentError("正式说明书文件缺失或完整性校验失败")
        return self._artifact_path(item["task_id"], item["docx_relative_path"]).read_bytes()

    def preview(self, job_id: str, version: int) -> dict:
        document = self.get(job_id, version)
        return {
            "document": document,
            "sections": self._sections(job_id),
            "figures": self._figures(job_id),
            "screenshots": self._screenshots(job_id),
        }

    def persist_toc_page_results(self, job_id: str, version: int,
                                 page_numbers: dict) -> dict:
        """Persist rendered PAGEREF results so exported DOCX opens with current pages."""
        document = self.get(job_id, version)
        path = self._artifact_path(document["task_id"], document["docx_relative_path"])
        if not page_numbers:
            return document
        with zipfile.ZipFile(path, "r") as source:
            xml = source.read("word/document.xml").decode("utf-8")
            pattern = re.compile(
                r'(<w:instrText[^>]*>PAGEREF manual_section_(\d+) \\h</w:instrText>'
                r'<w:fldChar w:fldCharType="separate"/><w:t>)(.*?)(</w:t>)'
            )

            def replace(match):
                number = page_numbers.get(int(match.group(2)))
                return match.group(0) if number is None else (
                    match.group(1) + str(number) + match.group(4)
                )

            updated_xml, count = pattern.subn(replace, xml)
            if not count or updated_xml == xml:
                return document
            with tempfile.NamedTemporaryFile(
                prefix=path.stem + "-toc-", suffix=".docx",
                dir=str(path.parent), delete=False,
            ) as handle:
                temporary_path = Path(handle.name)
            try:
                with zipfile.ZipFile(temporary_path, "w") as target:
                    for item in source.infolist():
                        payload = (updated_xml.encode("utf-8")
                                   if item.filename == "word/document.xml"
                                   else source.read(item.filename))
                        target.writestr(item, payload)
                os.replace(temporary_path, path)
            finally:
                temporary_path.unlink(missing_ok=True)
        sha256 = hashlib.sha256(path.read_bytes()).hexdigest()
        with self._database.connect() as connection:
            connection.execute(
                "UPDATE manual_document_artifacts SET sha256=? WHERE id=?",
                (sha256, document["id"]),
            )
        return self.get(job_id, version)

    def _context(self, job_id: str) -> dict:
        with self._database.connect() as connection:
            row = connection.execute(
                """SELECT j.id job_id, j.task_id, j.version job_version,
                ps.display_name project_name FROM manual_generation_jobs j
                JOIN tasks t ON t.id=j.task_id JOIN project_sources ps ON ps.id=t.source_id
                WHERE j.id=?""", (job_id,),
            ).fetchone()
            facts = connection.execute(
                """SELECT fact_key, value_json, status, created_at FROM facts
                WHERE task_id=(SELECT task_id FROM manual_generation_jobs WHERE id=?)
                AND fact_key IN ('project.name','project.version')
                AND status IN ('candidate','confirmed')
                ORDER BY CASE status WHEN 'confirmed' THEN 0 ELSE 1 END, created_at DESC""",
                (job_id,),
            ).fetchall()
        if row is None:
            raise ManualDocumentError("说明书生成任务不存在")
        result = dict(row)
        for fact in facts:
            key = fact["fact_key"]
            target = "software_name" if key == "project.name" else "software_version"
            if target not in result:
                result[target] = json.loads(fact["value_json"])
        result.setdefault("software_name", result["project_name"])
        result.setdefault("software_version", "V1.0")
        return result

    def _sections(self, job_id: str) -> list:
        with self._database.connect() as connection:
            rows = connection.execute(
                """SELECT * FROM manual_section_artifacts WHERE job_id=?
                ORDER BY ordinal""", (job_id,),
            ).fetchall()
        sections = []
        for row in rows:
            blocks = json.loads(row["content_json"])
            if row["section_key"] == "ui_operations":
                blocks = [block for block in blocks
                          if not self._internal_ui_delivery_block(block)]
            sections.append({
                "section_key": row["section_key"], "title": row["title"],
                "ordinal": row["ordinal"], "status": row["status"],
                "blocks": blocks,
                "evidence_refs": json.loads(row["evidence_refs_json"]),
                "inference_notes": json.loads(row["inference_notes_json"]),
            })
        return sections

    @classmethod
    def _internal_ui_delivery_block(cls, block: dict) -> bool:
        if block.get("type") == "subheading" and cls._generic_screenshot_group(
            block.get("title")
        ):
            return True
        text = str(block.get("text") or "")
        return any(marker in text for marker in (
            "本章节等待用户提供并确认真实截图证据",
            "其他正文章节可继续审阅与导出",
            "本章节仅标记为源码推断版",
            "用户已确认本项目不适用用户界面截图",
        ))

    def _figures(self, job_id: str) -> list:
        with self._database.connect() as connection:
            section_rows = connection.execute(
                "SELECT section_key, content_json, figure_requests_json FROM manual_section_artifacts WHERE job_id=?",
                (job_id,),
            ).fetchall()
            rows = connection.execute(
                """SELECT * FROM manual_figure_artifacts WHERE job_id=?
                AND status IN ('rendered','verified') ORDER BY section_key, figure_key""",
                (job_id,),
            ).fetchall()
        requested_keys = {
            item.get("figure_key") for section in section_rows
            for item in json.loads(section["figure_requests_json"] or "[]")
            if item.get("figure_key")
        }
        for section in section_rows:
            for block in json.loads(section["content_json"] or "[]"):
                if block.get("type") == "figure_request" and block.get("figure_key"):
                    requested_keys.add(block["figure_key"])
        request_types = {
            (section["section_key"], item.get("figure_type"))
            for section in section_rows
            for item in json.loads(section["figure_requests_json"] or "[]")
        }
        request_types.update(
            (section["section_key"], block.get("figure_type"))
            for section in section_rows
            for block in json.loads(section["content_json"] or "[]")
            if block.get("type") == "figure_request"
        )
        section_keys = {section["section_key"] for section in section_rows}
        if "architecture" in section_keys and ("architecture", "architecture") not in request_types:
            requested_keys.update({"system_architecture", "system-architecture"})
        if "modules" in section_keys and ("modules", "module") not in request_types:
            requested_keys.update({"module_collaboration", "module-collaboration"})
        if requested_keys:
            rows = [row for row in rows if row["figure_key"] in requested_keys]
        return [{
            "figure_key": row["figure_key"], "section_key": row["section_key"],
            "title": row["title"], "figure_type": row["figure_type"],
            "status": row["status"], "png_relative_path": row["png_relative_path"],
            "svg_relative_path": row["svg_relative_path"],
            "drawio_relative_path": row["drawio_relative_path"],
            "qa": json.loads(row["qa_json"]),
        } for row in rows]

    def _screenshots(self, job_id: str) -> list:
        with self._database.connect() as connection:
            evidence_rows = connection.execute(
                """SELECT a.id asset_id,a.asset_key,a.title,a.sensitive_status,
                r.image_relative_path,i.interpretation_json,i.version interpretation_version,
                ref.group_key,ref.group_title,ref.sort_order,i.id interpretation_revision_id
                FROM manual_job_screenshot_refs ref
                JOIN manual_project_screenshot_assets a ON a.id=ref.asset_id
                JOIN manual_project_screenshot_revisions r ON r.id=ref.asset_revision_id
                JOIN manual_screenshot_interpretation_revisions i
                    ON i.id=ref.interpretation_revision_id
                WHERE ref.job_id=? ORDER BY ref.sort_order,a.created_at,ref.group_title""",
                (job_id,),
            ).fetchall()
            if evidence_rows:
                screenshots = [{
                    "screenshot_key": row["asset_key"], "section_key": "ui_operations",
                    "asset_id": row["asset_id"],
                    "title": row["title"], "source": "project_evidence",
                    "image_relative_path": row["image_relative_path"],
                    "description": self._interpretation_description(
                        json.loads(row["interpretation_json"] or "{}")
                    ),
                    "interpretation": json.loads(row["interpretation_json"] or "{}"),
                    "interpretation_revision_id": row["interpretation_revision_id"],
                    "interpretation_version": row["interpretation_version"],
                    "group_key": row["group_key"], "group_title": row["group_title"],
                    "sort_order": row["sort_order"],
                    "sensitive_status": row["sensitive_status"],
                } for row in evidence_rows]
                for item in screenshots:
                    if self._generic_screenshot_group(item.get("group_title")):
                        suggested = str(item["interpretation"].get("suggested_group") or "").strip()
                        item["group_title"] = suggested or "界面说明"
                return screenshots
            rows = connection.execute(
                """SELECT * FROM manual_screenshot_artifacts WHERE job_id=?
                AND archived_at IS NULL
                ORDER BY section_key, created_at""", (job_id,),
            ).fetchall()
        return [{
            "screenshot_key": row["screenshot_key"], "section_key": row["section_key"],
            "title": row["title"], "source": row["source"],
            "image_relative_path": row["image_relative_path"],
            "description": json.loads(row["description_json"]),
            "interpretation_revision_id": "legacy:" + row["screenshot_key"],
            "sensitive_status": "confirmed_safe",
        } for row in rows]

    @staticmethod
    def _generic_screenshot_group(value: object) -> bool:
        normalized = re.sub(r"[\s_-]+", "", str(value or "")).lower()
        return normalized in {
            "", "截屏", "截图", "界面截图", "页面截图", "screenshot", "screenshots",
            "未分组", "未分组页面",
        }

    @staticmethod
    def _interpretation_description(value: dict) -> dict:
        def joined(key):
            item = value.get(key, [])
            return "；".join(str(entry) for entry in item) if isinstance(item, list) else str(item or "")
        backend = joined("related_backend_actions")
        warnings = joined("warnings")
        recovery = str(value.get("failure_and_recovery") or "")
        if warnings:
            recovery = (recovery + "；分析警告：" + warnings).strip("；")
        return {
            "page_purpose": str(value.get("purpose") or ""),
            "entry_conditions": joined("entry_conditions"),
            "visible_regions": "；".join(filter(None, [joined("visible_regions"),
                                                       joined("key_controls")])),
            "typical_workflow": joined("workflow_steps"),
            "backend_interactions": backend,
            "result_validation_recovery": "；".join(filter(None, [
                str(value.get("success_state") or ""), recovery,
            ])),
        }

    def _next_version(self, job_id: str) -> int:
        with self._database.connect() as connection:
            return connection.execute(
                """SELECT COALESCE(MAX(version),0)+1 value
                FROM manual_document_artifacts WHERE job_id=?""", (job_id,),
            ).fetchone()["value"]

    def _start_step(self, job_id: str) -> str:
        now = utc_now()
        with self._database.connect() as connection:
            row = connection.execute(
                """SELECT id,status,attempt FROM manual_generation_steps WHERE job_id=?
                AND step_key='assemble_docx' ORDER BY attempt DESC LIMIT 1""", (job_id,),
            ).fetchone()
            if row is None:
                raise ManualDocumentError("说明书任务缺少 DOCX 装配阶段")
            if row["status"] == "running":
                raise ManualDocumentError("正式说明书正在装配，请勿重复提交")
            if row["status"] in {"completed", "completed_with_warnings"}:
                step_id = str(uuid4())
                connection.execute(
                    """INSERT INTO manual_generation_steps(id,job_id,step_key,status,attempt,
                    summary_json,started_at) VALUES (?,?,'assemble_docx','running',?,'{}',?)""",
                    (step_id, job_id, row["attempt"] + 1, now),
                )
            else:
                step_id = row["id"]
                connection.execute(
                    """UPDATE manual_generation_steps SET status='running',started_at=?,
                    finished_at=NULL,safe_error_message=NULL WHERE id=?""", (now, step_id),
                )
            connection.execute(
                """UPDATE manual_generation_jobs SET status='running',current_step='assemble_docx',
                updated_at=?,safe_error_message=NULL WHERE id=?""", (now, job_id),
            )
        return step_id

    def _finish_failure(self, job_id: str, step_id: str, error: Exception) -> None:
        now = utc_now()
        message = "DOCX 装配失败：{0}".format(type(error).__name__)
        with self._database.connect() as connection:
            connection.execute(
                """UPDATE manual_generation_steps SET status='failed',finished_at=?,
                safe_error_message=? WHERE id=?""", (now, message, step_id),
            )
            connection.execute(
                """UPDATE manual_generation_jobs SET status='failed',current_step='assemble_docx',
                updated_at=?,safe_error_message=? WHERE id=?""", (now, message, job_id),
            )

    def _document_dict(self, context: dict, row) -> dict:
        relative = row["docx_relative_path"]
        artifact = self._artifact_path(context["task_id"], relative, require=False)
        status, size = "missing", None
        if artifact and artifact.is_file():
            digest = hashlib.sha256(artifact.read_bytes()).hexdigest()
            status = "verified" if digest == row["sha256"] else "mismatch"
            size = artifact.stat().st_size
        with self._database.connect() as connection:
            latest = connection.execute(
                """SELECT MAX(changed_at) changed_at FROM (
                    SELECT MAX(updated_at) changed_at FROM manual_section_artifacts WHERE job_id=?
                    UNION ALL
                    SELECT MAX(updated_at) changed_at FROM manual_figure_artifacts WHERE job_id=?
                    UNION ALL
                    SELECT MAX(COALESCE(NULLIF(updated_at,''),created_at)) changed_at
                    FROM manual_screenshot_artifacts WHERE job_id=?
                    UNION ALL
                    SELECT MAX(a.updated_at) changed_at
                    FROM manual_job_screenshot_refs ref
                    JOIN manual_project_screenshot_assets a ON a.id=ref.asset_id
                    WHERE ref.job_id=?
                )""", (row["job_id"], row["job_id"], row["job_id"], row["job_id"]),
            ).fetchone()["changed_at"]
            latest_qa = connection.execute(
                """SELECT qa_version, policy_version, passed, created_at
                FROM manual_document_qa_runs WHERE document_artifact_id=?
                ORDER BY qa_version DESC LIMIT 1""", (row["id"],),
            ).fetchone()
        freshness = "outdated" if latest and latest > row["created_at"] else "current"
        assembly = json.loads(row["qa_json"])
        # Keep the persisted status as immutable history, but never present an
        # artifact produced or checked by an obsolete policy as deliverable.
        from .manual_qa import QA_POLICY_VERSION
        generator_version = assembly.get("generator_version")
        current_generator = generator_version == GENERATOR_VERSION
        current_policy = bool(
            latest_qa and latest_qa["policy_version"] == QA_POLICY_VERSION
        )
        if latest_qa is None:
            quality_status = "not_checked"
        elif not latest_qa["passed"]:
            quality_status = "failed"
        elif not current_generator or not current_policy:
            quality_status = "outdated"
        else:
            quality_status = "passed"
        quality = {
            "status": quality_status,
            "passed": None if latest_qa is None else bool(latest_qa["passed"]),
            "qa_version": None if latest_qa is None else latest_qa["qa_version"],
            "policy_version": None if latest_qa is None else latest_qa["policy_version"],
            "current_policy": current_policy,
            "generator_version": generator_version,
            "current_generator": current_generator,
            "checked_at": None if latest_qa is None else latest_qa["created_at"],
        }
        return {
            "id": row["id"], "job_id": row["job_id"], "task_id": context["task_id"],
            "version": row["version"], "status": row["status"],
            "project_name": str(context["software_name"]),
            "project_version": str(context["software_version"]),
            "filename": self.export_filename(context, assembly.get("document_kind")),
            "document_kind": assembly.get("document_kind", "formal_candidate"),
            "docx_relative_path": relative, "preview_pdf_relative_path": row["preview_pdf_relative_path"],
            "qa": assembly, "quality": quality, "sha256": row["sha256"],
            "integrity": {"status": status, "size_bytes": size},
            "freshness": {"status": freshness, "latest_asset_update": latest},
            "created_at": row["created_at"],
        }

    def _artifact_path(self, task_id: str, relative: str,
                       require: bool = True) -> Optional[Path]:
        task_root = (self._data_root / "tasks" / task_id).resolve()
        artifact = (task_root / relative).resolve()
        if task_root not in artifact.parents:
            if require:
                raise ManualDocumentError("正式说明书路径越过任务目录")
            return None
        return artifact

    @staticmethod
    def export_filename(context: dict, document_kind: Optional[str] = None) -> str:
        def clean(value: object) -> str:
            result = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "-", str(value)).strip(" .")
            return result or "项目"
        suffix = {"review_checkpoint": "-阶段审阅稿",
                  "formal_candidate": "-审阅稿",
                  "final_document": ""}.get(document_kind, "-审阅稿")
        return "{0}-{1}-软件说明书{2}.docx".format(
            clean(context["software_name"]), clean(context["software_version"]), suffix
        )


class FormalManualBuilder:
    """Word-native A4 technical manual using the compact reference guide preset."""

    def build(self, output: Path, context: dict, sections: list, figures: list,
              screenshots: list, task_root: Path) -> dict:
        document = Document()
        self._configure_document(document, context)
        self._cover(document, context)
        document.add_page_break()
        self._toc(document, context, sections)
        document.add_page_break()
        figure_map = {item["figure_key"]: item for item in figures}
        figures_by_section = {}
        for item in figures:
            figures_by_section.setdefault(item["section_key"], []).append(item)
        screenshots_by_section = {}
        for item in screenshots:
            screenshots_by_section.setdefault(item["section_key"], []).append(item)
        inserted_figures, inserted_screenshots, warnings = 0, 0, []
        table_number = 0
        for fallback_index, section in enumerate(sections, 1):
            index = int(section.get("ordinal") or fallback_index)
            inserted_section_figures = set()
            subsection_index = 0
            compact_section = section.get("section_key") in {
                "ui_operations", "testing_summary",
            }
            existing_subheadings = {
                str(block.get("title") or "").strip()
                for block in section.get("blocks", []) if block.get("type") == "subheading"
            }
            screenshot_subheadings = {}
            for block in section.get("blocks", []):
                if block.get("type") != "subheading":
                    continue
                title = str(block.get("title") or "").strip()
                for evidence_ref in block.get("evidence_refs", []):
                    match = re.match(r"screenshot:([^:]+):", str(evidence_ref))
                    if title and match:
                        screenshot_subheadings[match.group(1)] = title
            heading = document.add_heading("{0}  {1}".format(index, section["title"]), level=1)
            self._bookmark(heading, "manual_section_{0}".format(index), index)
            # Let chapters flow continuously. A former special case forced the
            # UI chapter to a new page while its first screenshot heading also
            # inherited Word's heading pagination rules. LibreOffice could then
            # leave a header/footer-only page immediately before Chapter 7.
            # The first screenshot already stays with its picture, so no extra
            # chapter break is needed here.
            heading.paragraph_format.page_break_before = False
            for block in self._ordered_blocks(section):
                # Chapter 7 is assembled from the immutable adopted screenshot
                # snapshot below. Rendering its AI blocks here as well creates a
                # second, unsupported hierarchy before the actual screenshots.
                if section.get("section_key") == "ui_operations":
                    continue
                kind = block.get("type")
                if kind == "subheading":
                    subsection_index += 1
                    subheading = document.add_heading(
                        "{0}.{1}  {2}".format(index, subsection_index, block.get("title", "")),
                        level=2,
                    )
                    subheading.paragraph_format.page_break_before = False
                    subheading.paragraph_format.keep_with_next = True
                    if compact_section:
                        subheading.paragraph_format.space_before = Pt(8)
                        subheading.paragraph_format.space_after = Pt(4)
                elif kind == "paragraph":
                    self._body(
                        document, str(block.get("text", "")),
                        compact=compact_section,
                    )
                elif kind == "list":
                    lead = str(block.get("lead", "")).strip()
                    if lead:
                        self._body(document, lead)
                    for item in block.get("items", []):
                        self._bullet(document, str(item))
                elif kind == "table":
                    table_number += 1
                    self._table(document, block, table_number)
                elif kind == "figure_request":
                    figure = figure_map.get(block.get("figure_key"))
                    if figure and figure.get("png_relative_path"):
                        path = self._safe_asset(task_root, figure["png_relative_path"])
                        if path:
                            figure_number = inserted_figures + 1
                            self._picture(
                                document, path, figure["title"], "图", figure_number,
                                self._figure_height(figure.get("figure_type")),
                            )
                            purpose = str(block.get("purpose", "")).strip()
                            if purpose:
                                self._figure_note(document, purpose)
                            inserted_figures += 1
                            inserted_section_figures.add(figure["figure_key"])
                        else:
                            warnings.append("图表文件不可用：{0}".format(block.get("title", "未命名图")))
                    else:
                        warnings.append("图表尚未生成：{0}".format(block.get("title", "未命名图")))
            # Some evidence-derived figures are added by the deterministic
            # planner even when the model omitted a matching placeholder block.
            # They are still chapter assets and must not silently disappear from
            # the final DOCX.
            for figure in figures_by_section.get(section["section_key"], []):
                if figure["figure_key"] in inserted_section_figures:
                    continue
                path = self._safe_asset(task_root, figure.get("png_relative_path"))
                if not path:
                    warnings.append("图表文件不可用：{0}".format(figure["title"]))
                    continue
                self._picture(
                    document, path, figure["title"], "图", inserted_figures + 1,
                    self._figure_height(figure.get("figure_type")),
                )
                self._figure_note(document, "补充呈现本章的核心结构、职责边界与协作关系。")
                inserted_figures += 1
                inserted_section_figures.add(figure["figure_key"])
            section_screenshots = screenshots_by_section.get(section["section_key"], [])
            for screenshot_index, screenshot in enumerate(section_screenshots, 1):
                path = self._safe_asset(task_root, screenshot["image_relative_path"])
                if not path:
                    warnings.append("截图文件不可用：{0}".format(screenshot["title"]))
                    continue
                # A product screen and its explanation are one semantic unit.
                # Starting each screen on a dedicated page prevents a screenshot
                # from being left at the bottom of one page while its operational
                # notes spill onto the next page without context.
                title = self._clean_heading_number(
                    screenshot.get("interpretation", {}).get("page_title") or screenshot["title"]
                )
                screenshot_heading = document.add_heading(
                    "{0}.{1}  {2}".format(index, screenshot_index, title), level=2
                )
                # Keep the chapter title with its first real screen; later screens
                # each start on their own page for a stable screenshot/notes unit.
                screenshot_heading.paragraph_format.page_break_before = screenshot_index > 1
                screenshot_heading.paragraph_format.keep_with_next = True
                figure_number = inserted_figures + inserted_screenshots + 1
                if context.get("document_kind") != "final_document":
                    reference = document.add_paragraph(style="Normal")
                    reference.paragraph_format.keep_with_next = True
                    self._run(reference.add_run(
                        "对应操作界面见图 {0}，图中内容以已审核的真实截图为准。".format(
                            figure_number
                        )), 9.5, INK)
                self._picture(
                    document, path,
                    screenshot.get("interpretation", {}).get("suggested_caption") or
                    screenshot["title"], "图", figure_number, 5.0,
                )
                self._screenshot_description(
                    document, screenshot["description"],
                    final_document=context.get("document_kind") == "final_document",
                )
                inserted_screenshots += 1
        self._set_update_fields(document)
        document.save(str(output))
        self._embed_font(output)
        return {
            "generator_version": GENERATOR_VERSION,
            "design_preset": "compact_reference_guide",
            "named_override": "cn_software_manual_a4",
            "header_pattern": "editorial_cover",
            "section_count": len(sections), "figure_count": inserted_figures,
            "screenshot_count": inserted_screenshots, "warning_count": len(warnings),
            "warnings": warnings,
        }

    def _configure_document(self, document: Document, context: dict) -> None:
        section = document.sections[0]
        section.page_width, section.page_height = Cm(21), Cm(29.7)
        section.top_margin, section.bottom_margin = Cm(2.2), Cm(2.2)
        section.left_margin, section.right_margin = Cm(2.2), Cm(2.1)
        section.header_distance, section.footer_distance = Cm(1.25), Cm(1.25)
        styles = document.styles
        normal = styles["Normal"]
        self._style_font(normal, 10.5, INK)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(BODY_SPACE_AFTER_PT)
        # The A4 Chinese manual override uses a slightly tighter rhythm than the
        # Letter-sized preset. 1.25 made LibreOffice push short conclusions onto
        # an otherwise empty final page even though the deterministic companion
        # renderer reported a dense document.
        normal.paragraph_format.line_spacing = BODY_LINE_SPACING
        for name, size, color, before, after in (
            ("Heading 1", 16, ACCENT, 18, 10),
            ("Heading 2", 13, ACCENT, 14, 7),
            ("Heading 3", 12, ACCENT_DARK, 10, 5),
        ):
            style = styles[name]
            self._style_font(style, size, color, bold=True)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
        caption = styles["Caption"]
        self._style_font(caption, 9, MUTED)
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(8)
        caption.paragraph_format.keep_with_next = True
        if "Manual Screenshot Detail" not in styles:
            detail = styles.add_style("Manual Screenshot Detail", WD_STYLE_TYPE.PARAGRAPH)
        else:
            detail = styles["Manual Screenshot Detail"]
        self._style_font(detail, 10, INK)
        # Use one text grid for the Chapter 7 narrative and detail rows. Wrapped
        # lines remain aligned under the label/value rather than drifting left.
        detail.paragraph_format.left_indent = Pt(21)
        detail.paragraph_format.first_line_indent = Pt(0)
        detail.paragraph_format.space_after = Pt(BODY_SPACE_AFTER_PT)
        detail.paragraph_format.line_spacing = BODY_LINE_SPACING
        detail.paragraph_format.widow_control = True
        self._header_footer(section, context)
        document.core_properties.title = "{0} {1} 软件说明书".format(
            context["software_name"], context["software_version"]
        )
        document.core_properties.author = ""
        document.core_properties.last_modified_by = ""
        self._numbering(document)

    def _cover(self, document: Document, context: dict) -> None:
        for _ in range(5):
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(14)
        kicker = document.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(kicker.add_run("软件著作权登记材料"), 11, ACCENT, bold=True)
        kicker.paragraph_format.space_after = Pt(22)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(title.add_run(str(context["software_name"])), 28, INK, bold=True)
        title.paragraph_format.space_after = Pt(10)
        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        is_checkpoint = context.get("document_kind") == "review_checkpoint"
        is_final = context.get("document_kind") == "final_document"
        subtitle_text = ("软件说明书 · 阶段审阅稿" if is_checkpoint else
                         "软件说明书" if is_final else "软件说明书 · 审阅稿")
        self._run(subtitle.add_run(subtitle_text), 20, ACCENT_DARK, bold=True)
        subtitle.paragraph_format.space_after = Pt(24)
        version = document.add_paragraph()
        version.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(version.add_run(str(context["software_version"])), 13, MUTED, bold=True)
        version.paragraph_format.space_after = Pt(100)
        note = document.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_text = ("仅供正文审阅；图表、界面截图与最终质量检查尚未完成"
                     if is_checkpoint else "技术设计与使用说明")
        self._run(note.add_run(note_text), 10.5, MUTED)

    def _toc(self, document: Document, context: dict, sections: list) -> None:
        title = document.add_heading("目录", level=1)
        title.paragraph_format.page_break_before = False
        if context.get("document_kind") != "final_document":
            note = document.add_paragraph()
            note.paragraph_format.space_after = Pt(14)
            if context.get("document_kind") == "review_checkpoint":
                toc_note = ("阶段审阅稿：当前仅包含已完成正文；图表、界面截图、正式装配与质量检查"
                            "均可能尚未完成。页码将在 Word 打开或导出时自动更新。")
            else:
                toc_note = "章节导航（页码将在 Word 打开或导出时自动更新）"
            self._run(note.add_run(toc_note), 9.5, MUTED)
        for fallback_index, section in enumerate(sections, 1):
            index = int(section.get("ordinal") or fallback_index)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.3)
            paragraph.paragraph_format.right_indent = Cm(0.3)
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Cm(15.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
            )
            self._run(
                paragraph.add_run("{0:02d}    {1}\t".format(index, section["title"])),
                11, INK, bold=True,
            )
            self._field(paragraph, "PAGEREF manual_section_{0} \\h".format(index))

    def _body(self, document: Document, text: str, *, compact: bool = False) -> None:
        text = text.strip()
        if not text:
            return
        paragraph = document.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(21)
        paragraph.paragraph_format.widow_control = True
        if compact:
            # Operation and closing chapters contain many short semantic groups.
            # A compact but still legible rhythm prevents paragraph tails or the
            # final conclusion from becoming isolated on an otherwise empty page.
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(2)
        self._run(paragraph.add_run(text), 10 if compact else 10.5, INK)

    @staticmethod
    def _ordered_blocks(section: dict) -> list:
        """Place overview diagrams before responsibility tables when appropriate."""
        blocks = list(section.get("blocks", []))
        if any(item.get("type") == "subheading" for item in blocks):
            return blocks
        if section.get("section_key") not in {
            "architecture", "modules", "data_interfaces",
        }:
            return blocks
        figure_blocks = [item for item in blocks if item.get("type") == "figure_request"]
        if not figure_blocks:
            return blocks
        without_figures = [item for item in blocks if item.get("type") != "figure_request"]
        table_index = next(
            (index for index, item in enumerate(without_figures)
             if item.get("type") == "table"),
            len(without_figures),
        )
        return without_figures[:table_index] + figure_blocks + without_figures[table_index:]

    def _bullet(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
        ppr = paragraph._p.get_or_add_pPr()
        numpr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numid = OxmlElement("w:numId")
        numid.set(qn("w:val"), str(self._bullet_num_id))
        numpr.extend([ilvl, numid])
        ppr.append(numpr)
        self._run(paragraph.add_run(text.strip()), 10.5, INK)

    def _table(self, document: Document, block: dict, number: int) -> None:
        headers = [str(item) for item in block.get("headers", [])]
        rows = [[str(cell) for cell in row] for row in block.get("rows", [])]
        if not headers or not rows:
            return
        if block.get("title"):
            caption = document.add_paragraph(style="Caption")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._run(
                caption.add_run("表 {0}  {1}".format(number, block["title"])),
                9, MUTED,
            )
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        widths = self._column_widths(headers, rows)
        header_row = table.rows[0]
        self._repeat_header(header_row)
        for index, value in enumerate(headers):
            self._cell(header_row.cells[index], value, bold=True, fill=TABLE_FILL,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for row in rows:
            cells = table.add_row().cells
            self._cant_split(table.rows[-1])
            for index, value in enumerate(row[:len(headers)]):
                align = WD_ALIGN_PARAGRAPH.CENTER if len(value) < 18 else WD_ALIGN_PARAGRAPH.LEFT
                self._cell(cells[index], value, align=align)
        self._table_geometry(table, widths)
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_before = Pt(4)
        spacer.paragraph_format.space_after = Pt(4)

    def _picture(self, document: Document, path: Path, title: str, prefix: str,
                 number: int, max_height_inches: float) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run()
        with Image.open(path) as image:
            width_px, height_px = image.size
        if width_px / max(height_px, 1) >= 1.25:
            max_height_inches = min(max_height_inches, 3.35)
        width_inches = 6.25
        height_inches = width_inches * height_px / max(width_px, 1)
        if height_inches > max_height_inches:
            height_inches = max_height_inches
            width_inches = height_inches * width_px / max(height_px, 1)
        run.add_picture(
            str(path), width=Inches(width_inches), height=Inches(height_inches)
        )
        caption = document.add_paragraph(style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.keep_with_next = True
        self._run(
            caption.add_run("{0} {1}  {2}".format(prefix, number, title)),
            9, MUTED,
        )

    @staticmethod
    def _figure_height(figure_type: object) -> float:
        """Reserve prose room around diagrams instead of making every figure a page."""
        return {
            "workflow": 4.15,
            "data_flow": 4.35,
            "deployment": 3.75,
            "architecture": 3.65,
            "module": 5.05,
        }.get(str(figure_type or ""), 4.6)

    def _screenshot_description(self, document: Document, description: dict,
                                *, final_document: bool = False) -> None:
        purpose = str(description.get("page_purpose", "")).strip()
        workflow = str(description.get("typical_workflow", "")).strip()
        if purpose:
            narrative = "该功能模块主要用于{0}".format(purpose.rstrip("。"))
            if workflow:
                narrative += "。用户可在该页面按照界面提供的操作入口完成{0}".format(
                    workflow.rstrip("。").replace("；", "、")
                )
            if not final_document:
                narrative += "。页面实际内容和操作范围以图中已审核的真实界面为准。"
            self._body(document, narrative, compact=True)
        labels = (
            ("page_purpose", "页面用途"), ("entry_conditions", "进入条件"),
            ("visible_regions", "可见区域与控件"), ("typical_workflow", "典型操作流程"),
            ("backend_interactions", "后台、接口与数据交互"),
            ("result_validation_recovery", "结果、校验与异常恢复"),
        )
        for key, label in labels:
            value = str(description.get(key, "")).strip()
            if not value:
                continue
            paragraph = document.add_paragraph(style="Manual Screenshot Detail")
            label_run = paragraph.add_run(label + "：")
            self._run(label_run, 10, INK, bold=True)
            self._run(paragraph.add_run(value), 10, INK)

    @staticmethod
    def _clean_heading_number(value: object) -> str:
        """Word owns chapter numbering; model/user titles must not repeat it."""
        title = str(value or "").strip()
        return re.sub(
            r"^(?:第\s*\d+\s*[章节]\s*)?(?:\d+(?:\.\d+)*[\.、．]?\s*)+", "", title
        ) or "界面说明"

    def _figure_note(self, document: Document, purpose: str) -> None:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Cm(0.35)
        paragraph.paragraph_format.right_indent = Cm(0.35)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.keep_together = True
        self._run(paragraph.add_run("图示说明："), 9.5, ACCENT_DARK, bold=True)
        self._run(paragraph.add_run(purpose.rstrip("。") + "。"), 9.5, INK)

    def _header_footer(self, section, context: dict) -> None:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._run(header.add_run("{0} · 软件说明书".format(context["software_name"])),
                  8.5, MUTED)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(footer.add_run("第 "), 8.5, MUTED)
        self._field(footer, "PAGE")
        self._run(footer.add_run(" 页"), 8.5, MUTED)

    def _numbering(self, document: Document) -> None:
        numbering = document.part.numbering_part.element
        ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
        abstract_id = max(ids or [0]) + 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), "bullet")
        text = OxmlElement("w:lvlText")
        text.set(qn("w:val"), "•")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "540")
        indent.set(qn("w:hanging"), "271")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, indent, spacing])
        level.extend([start, fmt, text, ppr])
        abstract.append(level)
        numbering.append(abstract)
        num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
        self._bullet_num_id = max(num_ids or [0]) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(self._bullet_num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)

    @staticmethod
    def _safe_asset(task_root: Path, relative: str) -> Optional[Path]:
        root = task_root.resolve()
        path = (root / relative).resolve()
        if root not in path.parents or not path.is_file():
            return None
        return path

    @staticmethod
    def _style_font(style, size: float, color: str, bold: bool = False) -> None:
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold

    @staticmethod
    def _run(run, size: float, color: str, bold: bool = False) -> None:
        run.font.name = FONT_NAME
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold

    @staticmethod
    def _column_widths(headers: list, rows: list) -> list:
        weights = []
        for index, header in enumerate(headers):
            values = [header] + [row[index] for row in rows if index < len(row)]
            weights.append(min(36, max(8, max(len(value) for value in values))))
        total = sum(weights)
        widths = [round(TABLE_WIDTH_DXA * weight / total) for weight in weights]
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)
        return widths

    def _table_geometry(self, table, widths: list) -> None:
        tblpr = table._tbl.tblPr
        width = tblpr.first_child_found_in("w:tblW")
        width.set(qn("w:type"), "dxa")
        width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
        indent = OxmlElement("w:tblInd")
        indent.set(qn("w:type"), "dxa")
        indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
        tblpr.append(indent)
        existing_layout = tblpr.find(qn("w:tblLayout"))
        if existing_layout is None:
            existing_layout = OxmlElement("w:tblLayout")
            tblpr.append(existing_layout)
        existing_layout.set(qn("w:type"), "fixed")
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for value in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(value))
            grid.append(col)
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                tcw = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tcw.set(qn("w:type"), "dxa")
                tcw.set(qn("w:w"), str(widths[index]))
                cell.width = Inches(widths[index] / 1440)
                self._cell_margins(cell)

    def _cell(self, cell, value: str, bold: bool = False, fill: Optional[str] = None,
              align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
        value = {
            "ANONYMOUS": "匿名访问", "USER": "普通用户", "ADMIN": "管理员",
        }.get(value.strip().upper(), value)
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        self._run(paragraph.add_run(value), 9.5, INK, bold=bold)
        if fill:
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), fill)
            cell._tc.get_or_add_tcPr().append(shade)
        self._cell_margins(cell)

    @staticmethod
    def _cell_margins(cell) -> None:
        tcpr = cell._tc.get_or_add_tcPr()
        margins = tcpr.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            tcpr.append(margins)
        for name, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
            node = margins.find(qn("w:" + name))
            if node is None:
                node = OxmlElement("w:" + name)
                margins.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    @staticmethod
    def _repeat_header(row) -> None:
        trpr = row._tr.get_or_add_trPr()
        value = OxmlElement("w:tblHeader")
        value.set(qn("w:val"), "true")
        trpr.append(value)

    @staticmethod
    def _cant_split(row) -> None:
        trpr = row._tr.get_or_add_trPr()
        trpr.append(OxmlElement("w:cantSplit"))

    @staticmethod
    def _field(paragraph, instruction: str, placeholder: str = "1") -> None:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        code = OxmlElement("w:instrText")
        code.set(qn("xml:space"), "preserve")
        code.text = instruction
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        text = OxmlElement("w:t")
        text.text = placeholder
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, code, separate, text, end])

    @staticmethod
    def _bookmark(paragraph, name: str, bookmark_id: int) -> None:
        """Wrap a chapter heading in a stable bookmark for TOC page references."""
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        paragraph._p.insert(0, start)
        paragraph._p.append(end)

    @staticmethod
    def _set_update_fields(document: Document) -> None:
        settings = document.settings.element
        update = settings.find(qn("w:updateFields"))
        if update is None:
            update = OxmlElement("w:updateFields")
            settings.append(update)
        update.set(qn("w:val"), "true")

    @staticmethod
    def _embed_font(document_path: Path) -> None:
        """Embed the bundled OFL CJK font for portable Word and PDF rendering."""
        if not FONT_FILE.is_file():
            raise ManualDocumentError("内置中文字体文件缺失")
        font_key = uuid4()
        font_key_text = "{{{0}}}".format(str(font_key).upper())
        relationship_id = "rId1"
        font_entry = (
            '<w:font w:name="{0}"><w:family w:val="swiss"/>'
            '<w:pitch w:val="variable"/><w:embedRegular r:id="{1}" '
            'w:fontKey="{2}"/></w:font>'
        ).format(FONT_NAME, relationship_id, font_key_text).encode("utf-8")
        font_bytes = bytearray(FONT_FILE.read_bytes())
        mask = font_key.bytes[::-1]
        for index in range(min(32, len(font_bytes))):
            font_bytes[index] ^= mask[index % 16]
        relationships = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="{0}" Type="http://schemas.openxmlformats.org/'
            'officeDocument/2006/relationships/font" '
            'Target="fonts/NotoSansCJKsc-Regular.odttf"/></Relationships>'
        ).format(relationship_id).encode("utf-8")
        descriptor, temporary_name = tempfile.mkstemp(
            prefix="manual-font-table-", suffix=".docx", dir=str(document_path.parent)
        )
        os.close(descriptor)
        temporary = Path(temporary_name)
        try:
            with zipfile.ZipFile(document_path, "r") as source, zipfile.ZipFile(
                temporary, "w", compression=zipfile.ZIP_DEFLATED
            ) as target:
                for item in source.infolist():
                    payload = source.read(item.filename)
                    if item.filename == "word/fontTable.xml":
                        marker = b"</w:fonts>"
                        if marker not in payload:
                            raise ManualDocumentError("DOCX 字体表结构异常")
                        payload = payload.replace(marker, font_entry + marker, 1)
                    elif item.filename == "[Content_Types].xml":
                        marker = b"</Types>"
                        declaration = (
                            '<Default Extension="odttf" ContentType="application/vnd.'
                            'openxmlformats-officedocument.obfuscatedFont"/>'
                        ).encode("utf-8")
                        payload = payload.replace(marker, declaration + marker, 1)
                    elif item.filename == "word/settings.xml":
                        marker = b"</w:settings>"
                        settings = b"<w:embedTrueTypeFonts/><w:saveSubsetFonts/>"
                        payload = payload.replace(marker, settings + marker, 1)
                    target.writestr(item, payload)
                target.writestr("word/_rels/fontTable.xml.rels", relationships)
                target.writestr("word/fonts/NotoSansCJKsc-Regular.odttf", bytes(font_bytes))
            os.replace(str(temporary), str(document_path))
        finally:
            temporary.unlink(missing_ok=True)

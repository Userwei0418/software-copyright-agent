from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable
import unicodedata

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from .font_assets import (
    DEFAULT_CJK_FAMILY,
    DEFAULT_SYMBOL_FAMILY,
    FontAsset,
    OpenTypeCmap,
    embed_font_in_docx,
)


GENERATOR_VERSION = "source-docx-v4"


class SourceDocumentError(ValueError):
    pass


@dataclass(frozen=True)
class SourceDocumentTemplate:
    preset: str = "compact_reference_guide"
    cover_pattern: str = "editorial_cover"
    page_size: str = "A4"
    code_font: str = "Courier New"
    east_asia_font: str = DEFAULT_CJK_FAMILY
    code_size_pt: float = 9.0
    code_line_height_pt: float = 14.2
    code_pages: int = 59
    lines_per_page: int = 50


class SourceDocumentBuilder:
    def __init__(self, template: SourceDocumentTemplate = None) -> None:
        self.template = template or SourceDocumentTemplate()
        self._cjk_coverage = OpenTypeCmap(FontAsset.bundled_cjk().path)
        self._symbol_coverage = OpenTypeCmap(FontAsset.bundled_symbols().path)

    def build(
        self,
        output_path: Path,
        software_name: str,
        version: str,
        pages: Iterable[dict],
    ) -> dict:
        pages = list(pages)
        self._validate_pages(pages)
        document = Document()
        document.core_properties.title = "{0} {1} 源程序代码".format(
            software_name, version
        )
        document.core_properties.subject = "软件源程序代码"
        document.core_properties.author = ""
        document.core_properties.comments = ""
        self._configure_styles(document)

        cover_section = document.sections[0]
        self._configure_section(cover_section)
        cover_section.different_first_page_header_footer = True
        self._add_cover(document, software_name, version)

        code_section = document.add_section(WD_SECTION.NEW_PAGE)
        self._configure_section(code_section)
        code_section.different_first_page_header_footer = False
        code_section.header.is_linked_to_previous = False
        code_section.footer.is_linked_to_previous = False
        self._set_header(code_section, software_name, version)
        self._set_footer(code_section)

        for page_index, page in enumerate(pages):
            for entry in page["entries"]:
                self._add_code_line(document, entry)
            if page_index < len(pages) - 1:
                document.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        cjk_font_summary = embed_font_in_docx(output_path, FontAsset.bundled_cjk())
        symbol_font_summary = embed_font_in_docx(
            output_path, FontAsset.bundled_symbols()
        )
        return {
            "total_pages_expected": 1 + len(pages),
            "code_pages": len(pages),
            "lines_per_page": self.template.lines_per_page,
            "code_lines": sum(page["line_count"] for page in pages),
            "template": asdict(self.template),
            "embedded_fonts": [cjk_font_summary, symbol_font_summary],
        }

    def _validate_pages(self, pages: list) -> None:
        if len(pages) != self.template.code_pages:
            raise SourceDocumentError(
                "Source document requires exactly {0} code pages; got {1}".format(
                    self.template.code_pages, len(pages)
                )
            )
        for index, page in enumerate(pages, start=1):
            if page.get("page_number") != index:
                raise SourceDocumentError("Code preview page order is invalid")
            if page.get("line_count") != self.template.lines_per_page:
                raise SourceDocumentError(
                    "Code page {0} requires exactly {1} lines".format(
                        index, self.template.lines_per_page
                    )
                )
            if len(page.get("entries", [])) != self.template.lines_per_page:
                raise SourceDocumentError("Code page entry count is invalid")

    def _configure_styles(self, document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor(0, 0, 0)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25

    @staticmethod
    def _configure_section(section) -> None:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(17)
        section.right_margin = Mm(16)
        section.bottom_margin = Mm(15)
        section.left_margin = Mm(16)
        section.header_distance = Mm(8)
        section.footer_distance = Mm(8)

    def _add_cover(self, document: Document, software_name: str, version: str) -> None:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(116)

        kicker = document.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kicker.paragraph_format.space_after = Pt(18)
        run = kicker.add_run("SOFTWARE SOURCE CODE")
        self._set_run_font(run, "Calibri", 10, RGBColor(157, 117, 35), bold=True)

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(10)
        run = title.add_run(software_name)
        self._set_run_font(run, self.template.east_asia_font, 26, RGBColor(32, 55, 72), bold=True)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(34)
        run = subtitle.add_run("{0}  源程序代码".format(version))
        self._set_run_font(run, self.template.east_asia_font, 15, RGBColor(43, 81, 99))

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Pt(4)
        run = meta.add_run("文档构成：封面 1 页 + 源代码正文 59 页")
        self._set_run_font(run, self.template.east_asia_font, 10, RGBColor(80, 80, 80))

    def _set_header(self, section, software_name: str, version: str) -> None:
        paragraph = section.header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("{0}  {1}  源程序代码".format(software_name, version))
        self._set_run_font(run, self.template.east_asia_font, 8, RGBColor(90, 90, 90))

    def _set_footer(self, section) -> None:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run("第 ")
        self._set_run_font(run, self.template.east_asia_font, 8, RGBColor(100, 100, 100))
        self._add_field(paragraph, "PAGE")
        run = paragraph.add_run(" 页 / 共 ")
        self._set_run_font(run, self.template.east_asia_font, 8, RGBColor(100, 100, 100))
        self._add_field(paragraph, "NUMPAGES")
        run = paragraph.add_run(" 页")
        self._set_run_font(run, self.template.east_asia_font, 8, RGBColor(100, 100, 100))

    def _add_code_line(self, document: Document, entry: dict) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(self.template.code_line_height_pt)
        paragraph.paragraph_format.keep_with_next = False
        paragraph.paragraph_format.keep_together = False
        is_header = entry.get("kind") == "file_header"
        color = RGBColor(31, 77, 120) if is_header else RGBColor(0, 0, 0)
        for text, symbol_fallback in self._font_spans(entry.get("text", "")):
            run = paragraph.add_run(text)
            family = DEFAULT_SYMBOL_FAMILY if symbol_fallback else self.template.code_font
            self._set_run_font(
                run,
                family,
                self.template.code_size_pt,
                color,
                bold=is_header,
                east_asia=(DEFAULT_SYMBOL_FAMILY if symbol_fallback else self.template.east_asia_font),
            )

    def _font_spans(self, text: str) -> list:
        if not text:
            return [("", False)]
        spans = []
        current = []
        current_fallback = None
        for character in text:
            codepoint = ord(character)
            fallback = (
                codepoint > 127
                and not self._cjk_coverage.contains(codepoint)
                and (self._symbol_coverage.contains(codepoint)
                     or unicodedata.category(character).startswith("S"))
            )
            if current and fallback != current_fallback:
                spans.append(("".join(current), bool(current_fallback)))
                current = []
            current.append(character)
            current_fallback = fallback
        spans.append(("".join(current), bool(current_fallback)))
        return spans

    @staticmethod
    def _set_run_font(
        run,
        name: str,
        size: float,
        color: RGBColor,
        bold: bool = False,
        east_asia: str = None,
    ) -> None:
        run.font.name = name
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), name)
        fonts.set(qn("w:hAnsi"), name)
        fonts.set(qn("w:cs"), name)
        fonts.set(qn("w:eastAsia"), east_asia or name)

    @staticmethod
    def _add_field(paragraph, instruction: str) -> None:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        code = OxmlElement("w:instrText")
        code.set(qn("xml:space"), "preserve")
        code.text = instruction
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        text = OxmlElement("w:t")
        text.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend([begin, code, separate, text, end])
